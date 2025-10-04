from fastapi import FastAPI, Request, Form, HTTPException
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware
from dotenv import load_dotenv
import os
import logging
import sqlite3
import hashlib
import csv
import io
import tempfile
from io import BytesIO
import html
import hmac
import json
import re
import asyncio
from datetime import datetime, timedelta
from starlette.middleware.base import BaseHTTPMiddleware
from typing import Dict, List
from collections import defaultdict, deque
import time as _pytime
import time
import threading
from functools import wraps, lru_cache

from ollama_chat import get_response, get_response_with_messages
from typing import Dict, List, Callable, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
import traceback
import uuid

import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Tuple, BinaryIO
import magic

from fastapi import UploadFile, File
from fastapi.responses import FileResponse

from dataclasses import dataclass
from enum import Enum

import openai
import httpx
from abc import ABC, abstractmethod
try:
    import PyPDF2
    from docx import Document
    from PIL import Image
    import pytesseract
except ImportError as e:
    logger.warning(f"[FILES] Some file processing libraries missing: {e}")


def upgrade_database_for_threading():
    """Erweitert die Datenbank um Threading-Support"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Prüfen ob thread_id Spalte bereits existiert
        cursor.execute("PRAGMA table_info(chat_history)")
        columns = [column[1] for column in cursor.fetchall()]
        
        if 'thread_id' not in columns:
            # Thread-ID Spalte hinzufügen
            cursor.execute("ALTER TABLE chat_history ADD COLUMN thread_id TEXT DEFAULT 'default'")
            
            # Bestehende Nachrichten auf 'default' Thread setzen
            cursor.execute("UPDATE chat_history SET thread_id = 'default' WHERE thread_id IS NULL")
            
            logger.info("[THREADING] Database upgraded with thread_id column")
        
        # Threads-Tabelle erstellen
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS chat_threads (
                id TEXT PRIMARY KEY,
                username TEXT NOT NULL,
                title TEXT NOT NULL,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                message_count INTEGER DEFAULT 0,
                is_archived INTEGER DEFAULT 0,
                FOREIGN KEY (username) REFERENCES users (username)
            )
        """)
        
        # Default-Thread für alle User erstellen (falls nicht vorhanden)
        cursor.execute("""
            INSERT OR IGNORE INTO chat_threads (id, username, title, message_count)
            SELECT 'default', username, 'Haupt-Unterhaltung', 
                   (SELECT COUNT(*) FROM chat_history WHERE chat_history.username = users.username AND thread_id = 'default')
            FROM users
        """)
        
        conn.commit()
        logger.info("[THREADING] Threading system database setup complete")
        
    except Exception as e:
        logger.error(f"[THREADING] Database upgrade error: {e}")
        conn.rollback()
        raise
    finally:
        conn.close()

# ──────────────────────────────
# Thread Management Functions
# ──────────────────────────────

def create_new_thread(username: str, title: str = None) -> str:
    """Erstellt einen neuen Chat-Thread"""
    thread_id = str(uuid.uuid4())[:8]  # Kurze IDs für bessere UX
    
    if not title:
        # Auto-Titel basierend auf aktueller Zeit
        title = f"Chat {datetime.now().strftime('%d.%m. %H:%M')}"
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            INSERT INTO chat_threads (id, username, title, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?)
        """, (thread_id, username, title, datetime.now(), datetime.now()))
        
        conn.commit()
        logger.info(f"[THREADING] New thread created: {thread_id} for {username}")
        return thread_id
        
    except Exception as e:
        logger.error(f"[THREADING] Error creating thread: {e}")
        conn.rollback()
        raise
    finally:
        conn.close()

def get_user_threads(username: str) -> List[Dict]:
    """Holt alle Threads eines Users"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT t.id, t.title, t.created_at, t.updated_at, t.message_count, t.is_archived,
                   (SELECT content FROM chat_history 
                    WHERE username = ? AND thread_id = t.id AND role = 'user' 
                    ORDER BY timestamp DESC LIMIT 1) as last_message
            FROM chat_threads t 
            WHERE t.username = ? 
            ORDER BY t.updated_at DESC
        """, (username, username))
        
        rows = cursor.fetchall()
        
        threads = []
        for row in rows:
            threads.append({
                'id': row[0],
                'title': row[1],
                'created_at': datetime.fromisoformat(row[2]) if row[2] else datetime.now(),
                'updated_at': datetime.fromisoformat(row[3]) if row[3] else datetime.now(),
                'message_count': row[4] or 0,
                'is_archived': bool(row[5]),
                'last_message': row[6] or "Keine Nachrichten"
            })
        
        return threads
        
    except Exception as e:
        logger.error(f"[THREADING] Error getting threads for {username}: {e}")
        return []
    finally:
        conn.close()

def get_thread_history(username: str, thread_id: str) -> List[Dict]:
    """Holt Chat-Verlauf für einen bestimmten Thread"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT role, content, timestamp FROM chat_history 
            WHERE username = ? AND thread_id = ? 
            ORDER BY timestamp ASC
        """, (username, thread_id))
        
        rows = cursor.fetchall()
        return [{"role": row[0], "content": row[1], "timestamp": row[2]} for row in rows]
        
    except Exception as e:
        logger.error(f"[THREADING] Error getting thread history: {e}")
        return []
    finally:
        conn.close()

def save_message_to_thread(username: str, thread_id: str, role: str, content: str):
    """Speichert Nachricht in einem bestimmten Thread"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Nachricht speichern
        cursor.execute("""
            INSERT INTO chat_history (username, role, content, thread_id) 
            VALUES (?, ?, ?, ?)
        """, (username, role, content, thread_id))
        
        # Thread-Metadaten aktualisieren
        cursor.execute("""
            UPDATE chat_threads 
            SET updated_at = ?, message_count = message_count + 1
            WHERE id = ? AND username = ?
        """, (datetime.now(), thread_id, username))
        
        conn.commit()
        
    except Exception as e:
        logger.error(f"[THREADING] Error saving message to thread: {e}")
        conn.rollback()
        raise
    finally:
        conn.close()

def update_thread_title(username: str, thread_id: str, new_title: str) -> bool:
    """Ändert den Titel eines Threads"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            UPDATE chat_threads 
            SET title = ?, updated_at = ?
            WHERE id = ? AND username = ?
        """, (new_title, datetime.now(), thread_id, username))
        
        success = cursor.rowcount > 0
        conn.commit()
        
        if success:
            logger.info(f"[THREADING] Thread title updated: {thread_id} -> {new_title}")
        
        return success
        
    except Exception as e:
        logger.error(f"[THREADING] Error updating thread title: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def archive_thread(username: str, thread_id: str) -> bool:
    """Archiviert einen Thread (versteckt ihn aus der Hauptliste)"""
    if thread_id == 'default':
        return False  # Default-Thread kann nicht archiviert werden
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            UPDATE chat_threads 
            SET is_archived = 1, updated_at = ?
            WHERE id = ? AND username = ?
        """, (datetime.now(), thread_id, username))
        
        success = cursor.rowcount > 0
        conn.commit()
        return success
        
    except Exception as e:
        logger.error(f"[THREADING] Error archiving thread: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def delete_thread_completely(username: str, thread_id: str) -> bool:
    """Löscht einen Thread und alle seine Nachrichten komplett"""
    if thread_id == 'default':
        return False  # Default-Thread kann nicht gelöscht werden
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Erst alle Nachrichten löschen
        cursor.execute("DELETE FROM chat_history WHERE username = ? AND thread_id = ?", 
                      (username, thread_id))
        
        # Dann Thread löschen
        cursor.execute("DELETE FROM chat_threads WHERE id = ? AND username = ?", 
                      (thread_id, username))
        
        success = cursor.rowcount > 0
        conn.commit()
        
        if success:
            logger.info(f"[THREADING] Thread deleted completely: {thread_id}")
        
        return success
        
    except Exception as e:
        logger.error(f"[THREADING] Error deleting thread: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def get_thread_info(username: str, thread_id: str) -> Optional[Dict]:
    """Holt Informationen zu einem bestimmten Thread"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT id, title, created_at, updated_at, message_count, is_archived
            FROM chat_threads 
            WHERE id = ? AND username = ?
        """, (thread_id, username))
        
        row = cursor.fetchone()
        if not row:
            return None
        
        return {
            'id': row[0],
            'title': row[1],
            'created_at': datetime.fromisoformat(row[2]) if row[2] else datetime.now(),
            'updated_at': datetime.fromisoformat(row[3]) if row[3] else datetime.now(),
            'message_count': row[4] or 0,
            'is_archived': bool(row[5])
        }
        
    except Exception as e:
        logger.error(f"[THREADING] Error getting thread info: {e}")
        return None
    finally:
        conn.close()

# ──────────────────────────────
# Auto-Title Generation
# ──────────────────────────────

def generate_thread_title_from_first_message(first_message: str) -> str:
    """Generiert automatisch einen Titel basierend auf der ersten Nachricht"""
    # Erste 50 Zeichen nehmen und bei Wortgrenze abschneiden
    if len(first_message) <= 50:
        return first_message
    
    truncated = first_message[:47]
    last_space = truncated.rfind(' ')
    
    if last_space > 20:  # Mindestens 20 Zeichen
        return truncated[:last_space] + "..."
    else:
        return truncated + "..."

def update_thread_title_from_first_message(username: str, thread_id: str):
    """Aktualisiert Thread-Titel basierend auf erster User-Nachricht"""
    history = get_thread_history(username, thread_id)
    
    # Erste User-Nachricht finden
    first_user_message = None
    for msg in history:
        if msg['role'] == 'user':
            first_user_message = msg['content']
            break
    
    if first_user_message:
        new_title = generate_thread_title_from_first_message(first_user_message)
        update_thread_title(username, thread_id, new_title)

# ──────────────────────────────
# Threading Helper Functions
# ──────────────────────────────



class TaskStatus(Enum):
    PENDING = "pending"
    RUNNING = "running"
    COMPLETED = "completed"
    FAILED = "failed"
    CANCELLED = "cancelled"

@dataclass
class TaskInfo:
    """Info über einen Background-Task"""
    id: str
    name: str
    status: TaskStatus
    created_at: datetime
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    error_message: Optional[str] = None
    result: Any = None
    progress: float = 0.0
    metadata: Dict[str, Any] = field(default_factory=dict)

class BackgroundTaskManager:
    """
    Manager für Background-Tasks
    - Registriert und verwaltet Tasks
    - Überwacht Status und Ergebnisse
    - Bietet Admin-Interface für Monitoring
    """
    
    def __init__(self, max_concurrent_tasks: int = 5):
        self.tasks: Dict[str, TaskInfo] = {}
        self.running_tasks: Dict[str, asyncio.Task] = {}
        self.max_concurrent = max_concurrent_tasks
        self.semaphore = asyncio.Semaphore(max_concurrent_tasks)
        self.registered_task_functions: Dict[str, Callable] = {}
        
        # Periodische Tasks
        self.scheduled_tasks: Dict[str, Dict] = {}
        self.scheduler_running = False

    def register_task_function(self, name: str, func: Callable, description: str = ""):
        """Registriert eine Task-Funktion"""
        self.registered_task_functions[name] = {
            'function': func,
            'description': description
        }

    async def run_task(self, task_name: str, func: Callable, *args, **kwargs) -> str:
        """Startet einen Background-Task"""
        task_id = str(uuid.uuid4())[:8]
        
        # Task-Info erstellen
        task_info = TaskInfo(
            id=task_id,
            name=task_name,
            status=TaskStatus.PENDING,
            created_at=datetime.now()
        )
        
        self.tasks[task_id] = task_info
        
        # Task starten
        async_task = asyncio.create_task(self._execute_task(task_id, func, *args, **kwargs))
        self.running_tasks[task_id] = async_task
        
        return task_id

    async def _execute_task(self, task_id: str, func: Callable, *args, **kwargs):
        """Führt einen Task aus"""
        task_info = self.tasks[task_id]
        
        try:
            async with self.semaphore:  # Begrenzt gleichzeitige Tasks
                task_info.status = TaskStatus.RUNNING
                task_info.started_at = datetime.now()
                
                # Task ausführen
                if asyncio.iscoroutinefunction(func):
                    result = await func(*args, **kwargs)
                else:
                    result = func(*args, **kwargs)
                
                # Task erfolgreich
                task_info.status = TaskStatus.COMPLETED
                task_info.completed_at = datetime.now()
                task_info.result = result
                task_info.progress = 100.0
                
        except asyncio.CancelledError:
            task_info.status = TaskStatus.CANCELLED
            task_info.completed_at = datetime.now()
            
        except Exception as e:
            task_info.status = TaskStatus.FAILED
            task_info.completed_at = datetime.now()
            task_info.error_message = str(e)
            task_info.metadata['traceback'] = traceback.format_exc()
            
        finally:
            # Task aus running_tasks entfernen
            self.running_tasks.pop(task_id, None)

    def cancel_task(self, task_id: str) -> bool:
        """Bricht einen laufenden Task ab"""
        if task_id in self.running_tasks:
            self.running_tasks[task_id].cancel()
            return True
        return False

    def get_task_info(self, task_id: str) -> Optional[TaskInfo]:
        """Holt Task-Informationen"""
        return self.tasks.get(task_id)

    def get_all_tasks(self) -> List[TaskInfo]:
        """Alle Tasks mit Status"""
        return list(self.tasks.values())

    def get_running_tasks(self) -> List[TaskInfo]:
        """Nur laufende Tasks"""
        return [info for info in self.tasks.values() if info.status == TaskStatus.RUNNING]

    def cleanup_old_tasks(self, older_than_hours: int = 24):
        """Räumt alte Task-Einträge auf"""
        cutoff_time = datetime.now() - timedelta(hours=older_than_hours)
        
        tasks_to_remove = []
        for task_id, task_info in self.tasks.items():
            if (task_info.status in [TaskStatus.COMPLETED, TaskStatus.FAILED, TaskStatus.CANCELLED] and
                task_info.completed_at and task_info.completed_at < cutoff_time):
                tasks_to_remove.append(task_id)
        
        for task_id in tasks_to_remove:
            del self.tasks[task_id]
        
        return len(tasks_to_remove)

    def schedule_recurring_task(self, name: str, func: Callable, interval_seconds: int, 
                              run_immediately: bool = False, **kwargs):
        """Plant einen wiederkehrenden Task"""
        self.scheduled_tasks[name] = {
            'function': func,
            'interval': interval_seconds,
            'next_run': datetime.now() if run_immediately else datetime.now() + timedelta(seconds=interval_seconds),
            'kwargs': kwargs,
            'last_run': None,
            'run_count': 0
        }
        
        # Scheduler starten falls nicht schon aktiv
        if not self.scheduler_running:
            asyncio.create_task(self._scheduler_loop())

    async def _scheduler_loop(self):
        """Scheduler-Loop für wiederkehrende Tasks"""
        self.scheduler_running = True
        
        while self.scheduler_running:
            try:
                current_time = datetime.now()
                
                for name, schedule_info in self.scheduled_tasks.items():
                    if current_time >= schedule_info['next_run']:
                        # Task starten
                        await self.run_task(
                            f"{name} (scheduled)", 
                            schedule_info['function'],
                            **schedule_info['kwargs']
                        )
                        
                        # Nächsten Lauf planen
                        schedule_info['last_run'] = current_time
                        schedule_info['next_run'] = current_time + timedelta(seconds=schedule_info['interval'])
                        schedule_info['run_count'] += 1
                
                await asyncio.sleep(30)  # Alle 30 Sekunden prüfen
                
            except Exception as e:
                logger.error(f"Scheduler error: {e}")
                await asyncio.sleep(60)  # Bei Fehler länger warten

    def get_stats(self) -> Dict:
        """Task-Manager Statistiken"""
        total_tasks = len(self.tasks)
        running = len([t for t in self.tasks.values() if t.status == TaskStatus.RUNNING])
        completed = len([t for t in self.tasks.values() if t.status == TaskStatus.COMPLETED])
        failed = len([t for t in self.tasks.values() if t.status == TaskStatus.FAILED])
        cancelled = len([t for t in self.tasks.values() if t.status == TaskStatus.CANCELLED])
        
        return {
            'total_tasks': total_tasks,
            'running_tasks': running,
            'completed_tasks': completed,
            'failed_tasks': failed,
            'cancelled_tasks': cancelled,
            'max_concurrent': self.max_concurrent,
            'scheduled_tasks': len(self.scheduled_tasks),
            'scheduler_active': self.scheduler_running
        }

# ──────────────────────────────
# Spezifische Background Task Funktionen
# ──────────────────────────────

async def enhanced_cache_cleanup():
    """Erweiterte Cache-Bereinigung als Background-Task"""
    try:
        cleaned_count = app_cache.cleanup_expired()
        
        # Cache-Statistiken sammeln
        stats_before = app_cache.stats()
        
        # Weitere Bereinigung wenn nötig
        if stats_before['total_entries'] > 1000:
            # Alte User-Cache-Einträge löschen
            old_keys = [key for key in stats_before['cache_keys'] 
                       if 'user:' in key or 'user_persona:' in key]
            
            for key in old_keys[:100]:  # Maximal 100 auf einmal
                app_cache.delete(key)
        
        stats_after = app_cache.stats()
        
        return {
            'cleaned_expired': cleaned_count,
            'cache_entries_before': stats_before['total_entries'],
            'cache_entries_after': stats_after['total_entries'],
            'memory_freed_mb': stats_before['memory_usage_mb'] - stats_after['memory_usage_mb']
        }
        
    except Exception as e:
        logger.error(f"Cache cleanup task error: {e}")
        raise

async def database_maintenance():
    """Datenbank-Wartung als Background-Task"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        results = {}
        
        # VACUUM für Defragmentierung
        cursor.execute("VACUUM")
        results['vacuum_completed'] = True
        
        # Analyze für Statistiken
        cursor.execute("ANALYZE")
        results['analyze_completed'] = True
        
        # Alte Chat-Nachrichten bereinigen (älter als 90 Tage)
        cursor.execute("""
            DELETE FROM chat_history 
            WHERE timestamp < datetime('now', '-90 days')
        """)
        deleted_messages = cursor.rowcount
        results['deleted_old_messages'] = deleted_messages
        
        # Datenbankgröße ermitteln
        db_size = os.path.getsize(DB_PATH) if os.path.exists(DB_PATH) else 0
        results['database_size_mb'] = round(db_size / (1024 * 1024), 2)
        
        conn.commit()
        conn.close()
        
        return results
        
    except Exception as e:
        logger.error(f"Database maintenance task error: {e}")
        raise

async def rate_limit_cleanup():
    """Rate-Limit-Daten bereinigen"""
    try:
        limits = load_rate_limits()
        initial_users = len(limits)
        
        current_time = _pytime.time()
        day_ago = current_time - 86400  # 24 Stunden
        
        cleaned_users = 0
        for username, user_data in list(limits.items()):
            # Alte Nachrichten entfernen
            if 'messages' in user_data:
                old_count = len(user_data['messages'])
                user_data['messages'] = [
                    msg_time for msg_time in user_data['messages'] 
                    if msg_time > day_ago
                ]
                
            if 'daily_messages' in user_data:
                user_data['daily_messages'] = [
                    msg_time for msg_time in user_data['daily_messages'] 
                    if msg_time > day_ago
                ]
            
            # User ganz entfernen wenn keine aktuellen Nachrichten
            if (not user_data.get('messages') and 
                not user_data.get('daily_messages')):
                del limits[username]
                cleaned_users += 1
        
        save_rate_limits(limits)
        
        return {
            'initial_users': initial_users,
            'final_users': len(limits),
            'cleaned_users': cleaned_users
        }
        
    except Exception as e:
        logger.error(f"Rate limit cleanup task error: {e}")
        raise

async def system_health_check():
    """System-Gesundheitscheck"""
    try:
        results = {}
        
        # Speicherplatz prüfen
        import shutil
        disk_usage = shutil.disk_usage(".")
        results['disk_free_gb'] = round(disk_usage.free / (1024**3), 2)
        results['disk_used_percent'] = round(
            (disk_usage.used / disk_usage.total) * 100, 1
        )
        
        # Datenbank-Verbindung testen
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            cursor.execute("SELECT 1")
            conn.close()
            results['database_healthy'] = True
        except:
            results['database_healthy'] = False
        
        # Cache-Status
        cache_stats = app_cache.stats()
        results['cache_entries'] = cache_stats['total_entries']
        results['cache_memory_mb'] = cache_stats['memory_usage_mb']
        
        # Performance-Stats falls verfügbar
        if (hasattr(PerformanceMonitoringMiddleware, 'instance') and 
            PerformanceMonitoringMiddleware.instance):
            perf_stats = PerformanceMonitoringMiddleware.instance.get_stats()
            results['avg_response_time'] = perf_stats.get('average_response_time', 0)
            results['total_requests'] = perf_stats.get('total_requests', 0)
        
        return results
        
    except Exception as e:
        logger.error(f"Health check task error: {e}")
        raise

async def user_statistics_update():
    """User-Statistiken aktualisieren"""
    try:
        users = get_all_users()
        
        stats = {
            'total_users': len(users),
            'blocked_users': sum(1 for u in users.values() if u.get('blocked', False)),
            'admin_users': sum(1 for u in users.values() if u.get('is_admin', False)),
        }
        
        # Subscription-Tier-Verteilung
        tier_counts = {'free': 0, 'pro': 0, 'premium': 0}
        for user in users.values():
            tier = user.get('subscription_tier', 'free')
            if tier in tier_counts:
                tier_counts[tier] += 1
        
        stats['subscription_tiers'] = tier_counts
        
        # Chat-Aktivität (falls verfügbar)
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            
            # Nachrichten der letzten 24h
            cursor.execute("""
                SELECT COUNT(*) FROM chat_history 
                WHERE timestamp >= datetime('now', '-1 day')
            """)
            stats['messages_24h'] = cursor.fetchone()[0]
            
            # Aktive User (letzte 7 Tage)
            cursor.execute("""
                SELECT COUNT(DISTINCT username) FROM chat_history 
                WHERE timestamp >= datetime('now', '-7 days')
            """)
            stats['active_users_7d'] = cursor.fetchone()[0]
            
            conn.close()
        except:
            stats['messages_24h'] = 0
            stats['active_users_7d'] = 0
        
        return stats
        
    except Exception as e:
        logger.error(f"User statistics task error: {e}")
        raise

# ──────────────────────────────
# Task-Manager Setup und Registrierung
# ──────────────────────────────

def setup_background_tasks():
    """Registriert alle verfügbaren Background-Tasks"""
    
    # Task-Funktionen registrieren
    task_manager.register_task_function(
        "cache_cleanup", 
        enhanced_cache_cleanup,
        "Bereinigt abgelaufene Cache-Einträge und optimiert Speichernutzung"
    )
    
    task_manager.register_task_function(
        "database_maintenance",
        database_maintenance,
        "Führt VACUUM und ANALYZE aus, löscht alte Chat-Nachrichten"
    )
    
    task_manager.register_task_function(
        "rate_limit_cleanup",
        rate_limit_cleanup,
        "Bereinigt alte Rate-Limit-Daten"
    )
    
    task_manager.register_task_function(
        "system_health_check",
        system_health_check,
        "Überprüft System-Gesundheit (Speicher, DB, Performance)"
    )
    
    task_manager.register_task_function(
        "user_statistics",
        user_statistics_update,
        "Aktualisiert User- und Chat-Statistiken"
    )
    
    # Wiederkehrende Tasks planen
    task_manager.schedule_recurring_task(
        "cache_cleanup_scheduled",
        enhanced_cache_cleanup,
        interval_seconds=300,  # Alle 5 Minuten
        run_immediately=False
    )
    
    task_manager.schedule_recurring_task(
        "rate_limit_cleanup_scheduled", 
        rate_limit_cleanup,
        interval_seconds=3600,  # Stündlich
        run_immediately=False
    )
    
    task_manager.schedule_recurring_task(
        "health_check_scheduled",
        system_health_check, 
        interval_seconds=900,  # Alle 15 Minuten
        run_immediately=True
    )
    
    task_manager.schedule_recurring_task(
        "database_maintenance_scheduled",
        database_maintenance,
        interval_seconds=86400,  # Täglich
        run_immediately=False
    )
    
    setup_search_background_tasks()
    logger.info("[TASKS] Background-Task-System initialisiert")

# Alte cache_cleanup_background Funktion ersetzen
async def legacy_cache_cleanup_background():
    """Legacy-Funktion - wird ersetzt durch Task-Manager"""
    # Diese Funktion wird nicht mehr gebraucht
    # Der Task-Manager übernimmt das jetzt
    pass
# Globaler Task-Manager
task_manager = BackgroundTaskManager(max_concurrent_tasks=3)
# ──────────────────────────────
# Setup
# ──────────────────────────────
load_dotenv()
SECRET_KEY = os.getenv("SECRET_KEY", "supersecret")

app = FastAPI()

# Dann erst die Session Middleware:
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("app")

DB_PATH = "users.db"
CHAT_TABLE_CREATED = False
RATE_LIMIT_FILE = "rate_limits.json"
MESSAGES_PER_HOUR = 50
SESSION_TIMEOUT_MINUTES = 30
SESSION_TIMEOUT_SECONDS = SESSION_TIMEOUT_MINUTES * 60

# ──────────────────────────────
# Enhanced Caching System
# ──────────────────────────────

class SimpleCache:
    """Thread-safe Memory Cache mit TTL"""
    
    def __init__(self, default_ttl: int = 300):
        self.cache = {}
        self.ttl_data = {}
        self.lock = threading.Lock()
        self.default_ttl = default_ttl
        
    def get(self, key: str):
        with self.lock:
            if key in self.cache:
                if time.time() < self.ttl_data.get(key, 0):
                    return self.cache[key]
                else:
                    self.cache.pop(key, None)
                    self.ttl_data.pop(key, None)
            return None
    
    def set(self, key: str, value, ttl: int = None):
        ttl = ttl or self.default_ttl
        expire_time = time.time() + ttl
        
        with self.lock:
            self.cache[key] = value
            self.ttl_data[key] = expire_time
    
    def delete(self, key: str):
        with self.lock:
            self.cache.pop(key, None)
            self.ttl_data.pop(key, None)
    
    def clear(self):
        with self.lock:
            self.cache.clear()
            self.ttl_data.clear()
    
    def cleanup_expired(self):
        current_time = time.time()
        expired_keys = []
        
        with self.lock:
            for key, expire_time in self.ttl_data.items():
                if current_time >= expire_time:
                    expired_keys.append(key)
            
            for key in expired_keys:
                self.cache.pop(key, None)
                self.ttl_data.pop(key, None)
        
        return len(expired_keys)
    
    def stats(self):
        with self.lock:
            return {
                "total_entries": len(self.cache),
                "memory_usage_mb": round(
                    sum(len(str(v)) for v in self.cache.values()) / 1024 / 1024, 2
                ),
                "cache_keys": list(self.cache.keys())
            }

# Globaler Cache
app_cache = SimpleCache(default_ttl=300)

def cache_result(key_prefix: str, ttl: int = 300):
    """Decorator für Funktions-Caching"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            cache_key = f"{key_prefix}:{':'.join(map(str, args))}"
            if kwargs:
                cache_key += f":{hash(frozenset(kwargs.items()))}"
            
            cached = app_cache.get(cache_key)
            if cached is not None:
                return cached
            
            result = func(*args, **kwargs)
            
            if result is not None:
                app_cache.set(cache_key, result, ttl)
            
            return result
        return wrapper
    return decorator

def invalidate_user_cache(username: str):
    """User-bezogene Cache-Einträge löschen"""
    keys_to_delete = [
        f"user:{username}",
        f"user_persona:{username}",
        f"user_tier:{username}",
        f"available_personas:{username}"
    ]
    
    for key in keys_to_delete:
        app_cache.delete(key)
# ──────────────────────────────
# Threading Helper Functions 
# ──────────────────────────────

@cache_result("user_threads", ttl=300)
def get_user_threads_cached(username: str):
    """Cached Version von get_user_threads"""
    return get_user_threads(username)

def invalidate_thread_cache(username: str):
    """Thread-bezogene Cache-Einträge löschen"""
    app_cache.delete(f"user_threads:{username}")

# Session-Helper für aktuellen Thread
def get_current_thread_id(request) -> str:
    """Holt die aktuelle Thread-ID aus der Session"""
    return request.session.get("current_thread_id", "default")

def set_current_thread_id(request, thread_id: str):
    """Setzt die aktuelle Thread-ID in der Session"""
    request.session["current_thread_id"] = thread_id
    update_session_activity(request)

# ──────────────────────────────
# Database Extensions für Multi-AI Models
# ──────────────────────────────

def upgrade_database_for_models():
    """Erweitert die Datenbank um Multi-Model-Support"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # User-Model-Preference hinzufügen
        cursor.execute("PRAGMA table_info(users)")
        columns = [column[1] for column in cursor.fetchall()]
        
        if 'preferred_model' not in columns:
            cursor.execute("ALTER TABLE users ADD COLUMN preferred_model TEXT DEFAULT 'nexus'")
            logger.info("[MODELS] Added preferred_model column to users")
        
        # Chat-History um Model-Info erweitern
        cursor.execute("PRAGMA table_info(chat_history)")
        columns = [column[1] for column in cursor.fetchall()]
        
        if 'model_used' not in columns:
            cursor.execute("ALTER TABLE chat_history ADD COLUMN model_used TEXT DEFAULT 'nexus'")
            cursor.execute("ALTER TABLE chat_history ADD COLUMN tokens_used INTEGER DEFAULT 0")
            cursor.execute("ALTER TABLE chat_history ADD COLUMN response_time REAL DEFAULT 0.0")
            cursor.execute("ALTER TABLE chat_history ADD COLUMN cost REAL DEFAULT 0.0")
            logger.info("[MODELS] Added model tracking columns to chat_history")
        
        # Thread-Model-Preference
        cursor.execute("PRAGMA table_info(chat_threads)")
        columns = [column[1] for column in cursor.fetchall()]
        
        if 'preferred_model' not in columns:
            cursor.execute("ALTER TABLE chat_threads ADD COLUMN preferred_model TEXT DEFAULT NULL")
            logger.info("[MODELS] Added preferred_model column to chat_threads")
        
        conn.commit()
        logger.info("[MODELS] Database upgraded for multi-model support")
        
    except Exception as e:
        logger.error(f"[MODELS] Database upgrade error: {e}")
        conn.rollback()
        raise
    finally:
        conn.close()

# ──────────────────────────────
# Model Preference Management
# ──────────────────────────────
# ──────────────────────────────
# Model Preference Management + Cached wrappers
# ──────────────────────────────

# Verfügbare Modelle (cached)
@cache_result("available_models", ttl=300)
def get_available_models_for_user_cached(username: str) -> Dict[str, Any]:
    """
    Liefert die für einen User verfügbaren Modelle (cached).
    Greift auf AI_MODELS zurück, filtert nach Subscription.
    """
    try:
        tier = get_user_subscription_tier(username)
        tier_config = SUBSCRIPTION_TIERS.get(tier, SUBSCRIPTION_TIERS["free"])  # falls du es brauchst

        available_models = {}
        for model_id, model in AI_MODELS.items():
            if not getattr(model, "is_active", True):
                continue
            # Nur Modelle zulassen, die zum Tier passen (anpassen falls deine Logik anders ist)
            allowed_tier = getattr(model, "subscription_tier", "free")
            if allowed_tier not in [tier, "free"]:
                continue
            available_models[model_id] = model

        return available_models
    except Exception as e:
        logger.error(f"[MODELS] Error getting available models for {username}: {e}")
        return {}


def get_user_preferred_model(username: str) -> str:
    """Holt bevorzugtes Model des Users"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT preferred_model FROM users WHERE username = ?", (username,))
        result = cursor.fetchone()
        return result[0] if result and result[0] else "nexus"
    except sqlite3.OperationalError:
        # Spalte existiert noch nicht
        return "nexus"
    finally:
        conn.close()


def set_user_preferred_model(username: str, model_id: str):
    """Setzt bevorzugtes Model für User"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("UPDATE users SET preferred_model = ? WHERE username = ?", (model_id, username))
        conn.commit()
        logger.info(f"[MODELS] User {username} preferred model set to {model_id}")
        # Cache sofort invalidieren, damit UI direkt das neue Model sieht
        invalidate_model_cache(username)
    except Exception as e:
        logger.error(f"[MODELS] Error setting preferred model: {e}")
        conn.rollback()
    finally:
        conn.close()


def get_thread_preferred_model(username: str, thread_id: str) -> Optional[str]:
    """Holt Thread-spezifisches Model"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("""
            SELECT preferred_model FROM chat_threads 
            WHERE id = ? AND username = ?
        """, (thread_id, username))
        result = cursor.fetchone()
        return result[0] if result and result[0] else None
    except sqlite3.OperationalError:
        return None
    finally:
        conn.close()


def set_thread_preferred_model(username: str, thread_id: str, model_id: str):
    """Setzt Thread-spezifisches Model"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        cursor.execute("""
            UPDATE chat_threads SET preferred_model = ? 
            WHERE id = ? AND username = ?
        """, (model_id, thread_id, username))
        conn.commit()
        # Optional: auch hier Cache leeren, falls du thread-spezifisch cachen willst
        invalidate_model_cache(username)
    except Exception as e:
        logger.error(f"[MODELS] Error setting thread model: {e}")
        conn.rollback()
    finally:
        conn.close()


# ──────────────────────────────
# Cached wrappers & Invalidators (Models)
# ──────────────────────────────

@cache_result("user_preferred_model", ttl=300)
def get_user_preferred_model_cached(username: str) -> str:
    """Gecachte Variante von get_user_preferred_model."""
    return get_user_preferred_model(username)


def invalidate_model_cache(username: str):
    """Löscht Model-bezogene Cache-Keys für den User."""
    # verfügbare Modelle
    app_cache.delete(f"available_models:{username}")
    # bevorzugtes Nutzer-Modell
    app_cache.delete(f"user_preferred_model:{username}")

# ──────────────────────────────
# Model Usage Statistics
# ──────────────────────────────

def get_user_model_stats(username: str) -> Dict[str, Any]:
    """Liefert Nutzungs- und Kostenstatistiken pro Model für einen User."""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    try:
        # Daily Kosten (letzte 30 Tage)
        cursor.execute("""
            SELECT DATE(timestamp) as day,
                   SUM(cost) as total_cost,
                   COUNT(*) as requests
            FROM model_usage
            WHERE username = ? AND timestamp >= datetime('now', '-30 days')
            GROUP BY DATE(timestamp)
            ORDER BY day DESC
        """, (username,))
        daily = cursor.fetchall()

        # Gesamtkosten pro Modell
        cursor.execute("""
            SELECT model_name,
                   SUM(cost) as total_cost,
                   COUNT(*) as requests
            FROM model_usage
            WHERE username = ?
            GROUP BY model_name
            ORDER BY total_cost DESC
        """, (username,))
        totals = cursor.fetchall()

        return {
            "daily": [
                {"date": row[0], "cost": round(row[1] or 0, 4), "requests": row[2]}
                for row in daily
            ],
            "totals": [
                {"model": row[0], "cost": round(row[1] or 0, 4), "requests": row[2]}
                for row in totals
            ],
            "total_cost": round(sum(row[1] or 0 for row in totals), 4)
        }
    except sqlite3.OperationalError:
        # Tabelle model_usage existiert evtl. noch nicht
        return {"daily": [], "totals": [], "total_cost": 0}
    finally:
        conn.close()


# ──────────────────────────────
# Enhanced Chat Function mit Multi-Model-Support
# ──────────────────────────────

def save_message_to_thread_with_model(username: str, thread_id: str, role: str, 
                                     content: str, model_info: Dict = None):
    """Erweiterte Version mit Model-Tracking"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            INSERT INTO chat_history 
            (username, role, content, thread_id, model_used, tokens_used, response_time, cost) 
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            username, role, content, thread_id,
            model_info.get('model_used', 'nexus') if model_info else 'nexus',
            model_info.get('tokens_used', 0) if model_info else 0,
            model_info.get('duration', 0.0) if model_info else 0.0,
            model_info.get('cost', 0.0) if model_info else 0.0
        ))
        
        # Thread-Metadaten aktualisieren
        cursor.execute("""
            UPDATE chat_threads 
            SET updated_at = ?, message_count = message_count + 1
            WHERE id = ? AND username = ?
        """, (datetime.now(), thread_id, username))
        
        conn.commit()
        
    except sqlite3.OperationalError:
        # Fallback für alte DB-Schema
        cursor.execute("""
            INSERT INTO chat_history (username, role, content, thread_id) 
            VALUES (?, ?, ?, ?)
        """, (username, role, content, thread_id))
        conn.commit()
    except Exception as e:
        logger.error(f"[MODELS] Error saving message with model info: {e}")
        conn.rollback()
        raise
    finally:
        conn.close()

def get_model_context_for_thread(username: str, thread_id: str) -> str:
    """Bestimmt welches Model für Thread verwendet werden soll"""
    # 1. Thread-spezifisches Model
    thread_model = get_thread_preferred_model(username, thread_id)
    if thread_model:
        return thread_model
    
    # 2. User-bevorzugtes Model
    user_model = get_user_preferred_model(username)
    if user_model:
        return user_model
    
    # 3. Default
    return "nexus"

async def get_ai_response_with_model(current_message: str, chat_history: list,
                                   username: str, thread_id: str, persona: str = "standard",
                                   model_override: str = None) -> Dict:
    """
    Erweiterte KI-Response-Funktion mit Multi-Model-Support
    """
    # Model bestimmen
    model_to_use = model_override or get_model_context_for_thread(username, thread_id)

    # Nachrichten vorbereiten
    messages = []

    # System-Prompt basierend auf Persona + Model
    persona_config = PERSONAS.get(persona, PERSONAS["standard"])
    model_config = AI_MODELS.get(model_to_use, AI_MODELS["nexus"])
    system_prompt = f"{persona_config['system_prompt']}\n\nDu bist {model_config.display_name}: {model_config.personality}"

    messages.append({"role": "system", "content": system_prompt})

    # Chat-Historie (letzte 20 Nachrichten)
    for msg in chat_history[-20:]:
        if msg["role"] in ["user", "assistant"]:
            messages.append({"role": msg["role"], "content": msg["content"]})

    # aktuelle Nachricht
    messages.append({"role": "user", "content": current_message})

    # Response vom Model-Manager holen
    try:
        t0 = _pytime.time()
        result = await model_manager.generate_response(
            model_id=model_to_use,
            messages=messages,
            username=username,
            temperature=0.7,
            max_tokens=2000
        )
        duration = _pytime.time() - t0

        # Ergebnis normalisieren
        if isinstance(result, dict):
            content = result.get("content", "")
            tokens_used = result.get("tokens_used", 0)
            cost = result.get("cost", 0.0)
        else:
            content = str(result)
            tokens_used = 0
            cost = 0.0

        return {
            "content": content,
            "model_info": {
                "model_used": model_to_use,
                "provider": getattr(model_config, "provider", None).value if hasattr(model_config, "provider") else "unknown",
                "tokens_used": tokens_used,
                "duration": duration,
                "cost": cost,
                "is_fallback": False
            }
        }

    except Exception as e:
        logger.error(f"[MODELS] generate_response failed for {model_to_use}: {e}")

        # Fallback – versuch über get_response_with_messages
        try:
            t1 = _pytime.time()
            fallback_content = get_response_with_messages(messages)
            duration_fb = _pytime.time() - t1
        except Exception as inner:
            logger.error(f"[MODELS] Fallback failed: {inner}")
            fallback_content = "Entschuldige, es ist ein Fehler aufgetreten."
            duration_fb = 0.0

        return {
            "content": fallback_content,
            "model_info": {
                "model_used": "fallback_local",
                "provider": "local",
                "tokens_used": 0,
                "duration": duration_fb,
                "cost": 0.0,
                "is_fallback": True
            }
        }

        
        

# ──────────────────────────────
# Subscription Tiers Definition
# ──────────────────────────────
SUBSCRIPTION_TIERS = {
    "free": {
        "name": "Free",
        "max_messages_per_hour": 10,
        "max_messages_per_day": 50,
        "available_personas": ["standard", "freundlich"],
        "features": ["Basis-Chat", "2 Personas"]
    },
    "pro": {
        "name": "Pro", 
        "max_messages_per_hour": 100,
        "max_messages_per_day": 500,
        "available_personas": ["standard", "freundlich", "lustig", "professionell"],
        "features": ["Erweiterte Personas", "Mehr Nachrichten", "Priorität"]
    },
    "premium": {
        "name": "Premium",
        "max_messages_per_hour": -1,  # Unlimited
        "max_messages_per_day": -1,   # Unlimited
        "available_personas": ["standard", "freundlich", "lustig", "professionell", "lehrerin", "kreativ", "analyst", "therapeut"],
        "features": ["Alle Personas", "Unbegrenzte Nachrichten", "Höchste Priorität", "Erweiterte Features"]
    }
}

# Enhanced Persona-Definitionen mit Tier-Zuordnung
PERSONAS = {
    "standard": {
        "name": "Standard Assistent",
        "emoji": "🤖",
        "system_prompt": "Du bist ein hilfsfreundlicher KI-Assistent. Antworte auf Deutsch und sei sachlich aber freundlich.",
        "tier": "free",
        "description": "Der klassische KI-Assistent für alltägliche Fragen"
    },
    "freundlich": {
        "name": "Freundlicher Helfer", 
        "emoji": "😊",
        "system_prompt": "Du bist ein sehr freundlicher und enthusiastischer Assistent. Verwende warme, ermutigende Worte und zeige echtes Interesse an den Fragen. Sei optimistisch und unterstützend.",
        "tier": "free",
        "description": "Besonders warmherzig und ermutigend"
    },
    "lustig": {
        "name": "Comedy Bot",
        "emoji": "😄", 
        "system_prompt": "Du bist ein humorvoller Assistent der gerne Witze macht und lustige Antworten gibt. Bleibe trotzdem hilfreich, aber bringe den User zum Lächeln. Verwende gelegentlich Wortwitz oder lustige Vergleiche.",
        "tier": "pro",
        "description": "Bringt Humor in jede Unterhaltung"
    },
    "professionell": {
        "name": "Business Experte",
        "emoji": "👔",
        "system_prompt": "Du bist ein professioneller Berater mit Expertise in Business und Technik. Antworte präzise, strukturiert und sachlich. Nutze Fachbegriffe angemessen und gib konkrete Handlungsempfehlungen.",
        "tier": "pro",
        "description": "Für geschäftliche und technische Anfragen"
    },
    "lehrerin": {
        "name": "Geduldige Lehrerin", 
        "emoji": "👩‍🏫",
        "system_prompt": "Du bist eine geduldige Lehrerin die komplexe Themen einfach erklärt. Baue Erklärungen schrittweise auf, verwende Beispiele und frage nach ob alles verstanden wurde. Ermutige zum Lernen.",
        "tier": "premium",
        "description": "Erklärt komplexe Themen verständlich"
    },
    "kreativ": {
        "name": "Kreativer Geist",
        "emoji": "🎨", 
        "system_prompt": "Du bist ein kreativer Assistent voller Ideen und Inspiration. Denke um die Ecke, schlage ungewöhnliche Lösungen vor und bringe künstlerische Perspektiven ein. Sei experimentierfreudig.",
        "tier": "premium",
        "description": "Für kreative Projekte und Inspiration"
    },
    "analyst": {
        "name": "Daten Analyst",
        "emoji": "📊",
        "system_prompt": "Du bist ein präziser Datenanalyst. Analysiere Informationen systematisch, identifiziere Muster und Trends, und präsentiere Erkenntnisse klar strukturiert mit Zahlen und Fakten.",
        "tier": "premium", 
        "description": "Spezialist für Datenanalyse und Statistiken"
    },
    "therapeut": {
        "name": "Empathischer Berater",
        "emoji": "🧘‍♀️",
        "system_prompt": "Du bist ein einfühlsamer Gesprächspartner der aktiv zuhört und unterstützt. Stelle durchdachte Fragen, biete verschiedene Perspektiven und hilfe dabei, Gedanken zu ordnen. Sei verständnisvoll aber nicht direktiv.",
        "tier": "premium",
        "description": "Für persönliche Gespräche und Reflexion"
    }
}

# ──────────────────────────────
# Performance Monitoring Middleware
# ──────────────────────────────

class PerformanceMonitoringMiddleware(BaseHTTPMiddleware):
    """
    Middleware für Performance-Monitoring
    - Misst Request-Zeiten
    - Sammelt Statistiken
    - Speichert langsame Requests
    """

    instance = None  # Statische Referenz auf die Instanz
    
    def __init__(self, app, slow_threshold: float = 2.0):
        super().__init__(app)
        self.slow_threshold = slow_threshold
        self.stats = {
            'total_requests': 0,
            'total_time': 0.0,
            'slow_requests': deque(maxlen=100),  # Letzte 100 langsame Requests
            'endpoint_stats': defaultdict(lambda: {'count': 0, 'total_time': 0.0}),
            'last_requests': deque(maxlen=50)  # Letzte 50 Requests
        }
        PerformanceMonitoringMiddleware.instance = self  # Instanz merken

    async def dispatch(self, request: Request, call_next):
        start_time = _pytime.time()

        # Request verarbeiten
        response = await call_next(request)

        # Performance messen
        process_time = _pytime.time() - start_time

        # Statistiken aktualisieren
        self._update_stats(request, process_time)

        # Performance-Header hinzufügen
        response.headers["X-Process-Time"] = str(round(process_time, 4))

        return response

    def _update_stats(self, request: Request, process_time: float):
        """Interne Statistiken aktualisieren"""
        path = request.url.path
        method = request.method
        endpoint = f"{method} {path}"

        # Global stats
        self.stats['total_requests'] += 1
        self.stats['total_time'] += process_time

        # Endpoint stats
        self.stats['endpoint_stats'][endpoint]['count'] += 1
        self.stats['endpoint_stats'][endpoint]['total_time'] += process_time

        # Langsame Requests tracken
        if process_time > self.slow_threshold:
            self.stats['slow_requests'].append({
                'timestamp': _pytime.time(),
                'endpoint': endpoint,
                'duration': round(process_time, 4),
                'user_agent': request.headers.get('user-agent', 'Unknown')[:100]
            })

        # Letzte Requests tracken
        self.stats['last_requests'].append({
            'timestamp': _pytime.time(),
            'endpoint': endpoint,
            'duration': round(process_time, 4),
            'status': 'slow' if process_time > self.slow_threshold else 'normal'
        })

    def get_stats(self) -> Dict:
        """Performance-Statistiken zurückgeben"""
        avg_time = (self.stats['total_time'] / self.stats['total_requests']
                   if self.stats['total_requests'] > 0 else 0)

        # Endpoint-Statistiken aufbereiten
        endpoint_stats = []
        for endpoint, data in self.stats['endpoint_stats'].items():
            avg_endpoint_time = data['total_time'] / data['count']
            endpoint_stats.append({
                'endpoint': endpoint,
                'count': data['count'],
                'avg_time': round(avg_endpoint_time, 4),
                'total_time': round(data['total_time'], 4)
            })

        # Nach durchschnittlicher Zeit sortieren
        endpoint_stats.sort(key=lambda x: x['avg_time'], reverse=True)

        return {
            'total_requests': self.stats['total_requests'],
            'average_response_time': round(avg_time, 4),
            'total_time': round(self.stats['total_time'], 4),
            'slow_requests_count': len(self.stats['slow_requests']),
            'slow_threshold': self.slow_threshold,
            'endpoint_stats': endpoint_stats[:10],  # Top 10 langsamste
            'recent_slow_requests': list(self.stats['slow_requests'])[-10:],  # Letzte 10 langsame
            'recent_requests': list(self.stats['last_requests'])[-20:]  # Letzte 20 allgemein
        }

# Middleware einbinden
app.add_middleware(PerformanceMonitoringMiddleware, slow_threshold=2.0)

# ──────────────────────────────
# Database Helper Functions
# ──────────────────────────────
def init_db():
    """Initialisiert die Datenbank mit allen nötigen Tabellen"""
    global CHAT_TABLE_CREATED
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Users-Tabelle (konsistent mit deinem reset.py)
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT NOT NULL,
            question TEXT NOT NULL,
            answer TEXT NOT NULL,
            is_admin INTEGER DEFAULT 0,
            is_blocked INTEGER DEFAULT 0,
            persona TEXT DEFAULT 'standard',
            subscription_tier TEXT DEFAULT 'free'
        )
    """)
    
    # Chat-Verlauf Tabelle
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (username) REFERENCES users (username)
        )
    """)
    
    # Admin-User erstellen falls nicht existiert
    cursor.execute("SELECT username FROM users WHERE username = ?", ("admin",))
    if not cursor.fetchone():
        admin_hash = hash_password("admin")
        cursor.execute("""
            INSERT INTO users (username, password, question, answer, is_admin, subscription_tier) 
            VALUES (?, ?, ?, ?, 1, 'premium')
        """, ("admin", admin_hash, "Default Admin Question", "admin"))
        logger.info("[INIT] Admin-User erstellt (admin/admin)")
    
    conn.commit()
    conn.close()
    CHAT_TABLE_CREATED = True
    logger.info("[INIT] Datenbank initialisiert")

def hash_password(password: str) -> str:
    """Passwort hashen"""
    return hashlib.sha256(password.encode()).hexdigest()

def render_markdown_simple(text: str) -> str:
    """
    Einfaches Markdown-Rendering für Chat-Nachrichten
    Unterstützt: Code-Blöcke, Inline-Code, Listen, Links, Fett/Kursiv
    """
    # HTML escaping für Sicherheit
    text = html.escape(text)
    
    # Code-Blöcke (```code```)
    text = re.sub(
        r'```(\w+)?\n?(.*?)```', 
        r'<div class="code-block"><div class="code-header">\1</div><pre><code>\2</code></pre></div>', 
        text, 
        flags=re.DOTALL
    )
    
    # Inline-Code (`code`)
    text = re.sub(r'`([^`]+)`', r'<code class="inline-code">\1</code>', text)
    
    # Fett (**text**)
    text = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', text)
    
    # Kursiv (*text*)
    text = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', text)
    
    # Listen (- item)
    lines = text.split('\n')
    in_list = False
    result_lines = []
    
    for line in lines:
        if line.strip().startswith('- '):
            if not in_list:
                result_lines.append('<ul class="chat-list">')
                in_list = True
            result_lines.append(f'<li>{line.strip()[2:]}</li>')
        else:
            if in_list:
                result_lines.append('</ul>')
                in_list = False
            result_lines.append(line)
    
    if in_list:
        result_lines.append('</ul>')
    
    text = '\n'.join(result_lines)
    
    # Links [text](url)
    text = re.sub(
        r'\[([^\]]+)\]\(([^)]+)\)', 
        r'<a href="\2" target="_blank" class="chat-link">\1</a>', 
        text
    )
    
    # Zeilenumbrüche
    text = text.replace('\n', '<br>')
    
    return text

def get_user(username: str) -> dict:
    """User aus DB holen"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE username = ?", (username,))
    row = cursor.fetchone()
    conn.close()
    
    if row:
        return {
            "username": row[0],
            "password": row[1],
            "question": row[2],
            "answer": row[3],
            "is_admin": bool(row[4]),
            "is_blocked": bool(row[5]),
            "persona": row[6] if len(row) > 6 else "standard",
            "subscription_tier": row[7] if len(row) > 7 else "free"
        }
    return None

def get_all_users() -> dict:
    """Alle User für Admin-Panel"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users")
    rows = cursor.fetchall()
    conn.close()
    
    users = {}
    for row in rows:
        users[row[0]] = {
            "password": row[1],
            "question": row[2],
            "answer": row[3],
            "is_admin": bool(row[4]),
            "blocked": bool(row[5]),  # für Template-Kompatibilität
            "subscription_tier": row[7] if len(row) > 7 else "free"
        }
    return users

def save_user(username: str, password: str, question: str, answer: str) -> bool:
    """Neuen User registrieren"""
    if get_user(username):
        return False  # User existiert bereits
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        password_hash = hash_password(password)
        cursor.execute("""
            INSERT INTO users (username, password, question, answer, subscription_tier) 
            VALUES (?, ?, ?, ?, 'free')
        """, (username, password_hash, question, answer))
        conn.commit()
        logger.info(f"[REGISTER] Neuer User: {username}")
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def check_login(username: str, password: str) -> bool:
    """Login überprüfen"""
    user = get_user(username)
    if user:
        return user["password"] == hash_password(password)
    return False

def verify_security_answer(username: str, answer: str) -> bool:
    """Sicherheitsantwort prüfen"""
    user = get_user(username)
    if user:
        return user["answer"] == answer
    return False

def reset_password(username: str, new_password: str):
    """Passwort zurücksetzen"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    password_hash = hash_password(new_password)
    cursor.execute("UPDATE users SET password = ? WHERE username = ?", 
                   (password_hash, username))
    conn.commit()
    conn.close()

def get_user_history(username: str) -> list:
    """Chat-Verlauf aus DB laden"""
    if not CHAT_TABLE_CREATED:
        return []
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        SELECT role, content, timestamp FROM chat_history 
        WHERE username = ? ORDER BY timestamp ASC
    """, (username,))
    rows = cursor.fetchall()
    conn.close()
    
    return [{"role": row[0], "content": row[1], "timestamp": row[2]} for row in rows]

# ──────────────────────────────
# Cached Database Functions  
# ──────────────────────────────

@cache_result("user", ttl=600)
def get_user_cached(username: str):
    """Cached Version von get_user()"""
    return get_user(username)

@cache_result("user_persona", ttl=300)
def get_user_persona_cached(username: str):
    """Cached Version von get_user_persona()"""
    return get_user_persona(username)

@cache_result("user_tier", ttl=600)
def get_user_subscription_tier_cached(username: str):
    """Cached Version von get_user_subscription_tier()"""
    return get_user_subscription_tier(username)

@cache_result("available_personas", ttl=600)
def get_available_personas_for_user_cached(username: str):
    """Cached Version von get_available_personas_for_user()"""
    return get_available_personas_for_user(username)

@cache_result("all_users", ttl=120)
def get_all_users_cached():
    """Cached Version von get_all_users()"""
    return get_all_users()

# ──────────────────────────────
# Erweiterte Funktionen mit Cache-Invalidation
# ──────────────────────────────

def save_user_persona_cached(username: str, persona: str):
    """save_user_persona mit Cache-Invalidierung"""
    save_user_persona(username, persona)
    invalidate_user_cache(username)

def search_chat_messages(username: str, query: str, thread_id: str = None, 
                        days_back: int = None, limit: int = 50) -> List[Dict]:
    """
    Durchsucht Chat-Nachrichten eines Users
    
    Args:
        username: Der User
        query: Suchbegriff
        thread_id: Optionaler Thread-Filter (None = alle Threads)
        days_back: Optionaler Zeitfilter (None = alle Zeit)
        limit: Max. Anzahl Ergebnisse
    
    Returns:
        Liste von Suchergebnissen mit Kontext
    """
    if not query.strip():
        return []
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Base Query
        sql = """
            SELECT h.id, h.role, h.content, h.timestamp, h.thread_id,
                   t.title as thread_title
            FROM chat_history h
            LEFT JOIN chat_threads t ON h.thread_id = t.id AND h.username = t.username
            WHERE h.username = ? AND h.content LIKE ?
        """
        params = [username, f"%{query}%"]
        
        # Thread-Filter
        if thread_id:
            sql += " AND h.thread_id = ?"
            params.append(thread_id)
        
        # Zeit-Filter
        if days_back:
            cutoff_date = datetime.now() - timedelta(days=days_back)
            sql += " AND h.timestamp >= ?"
            params.append(cutoff_date.isoformat())
        
        sql += " ORDER BY h.timestamp DESC LIMIT ?"
        params.append(limit)
        
        cursor.execute(sql, params)
        rows = cursor.fetchall()
        
        results = []
        for row in rows:
            # Kontext holen (vorherige und nächste Nachricht)
            context = get_message_context(username, row[0], row[4])  # message_id, thread_id
            
            results.append({
                'id': row[0],
                'role': row[1],
                'content': row[2],
                'timestamp': row[3],
                'thread_id': row[4],
                'thread_title': row[5] or 'Unbekannter Thread',
                'highlighted_content': highlight_search_term(row[2], query),
                'context': context
            })
        
        return results
        
    except Exception as e:
        logger.error(f"[SEARCH] Error searching messages for {username}: {e}")
        return []
    finally:
        conn.close()

def get_message_context(username: str, message_id: int, thread_id: str, 
                       context_size: int = 2) -> Dict:
    """
    Holt Kontext um eine bestimmte Nachricht
    
    Args:
        username: Der User
        message_id: ID der Nachricht
        thread_id: Thread der Nachricht
        context_size: Anzahl Nachrichten vor/nach der gefundenen
    
    Returns:
        Dict mit before/after Nachrichten
    """
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Nachrichten vor der gefundenen
        cursor.execute("""
            SELECT role, content, timestamp 
            FROM chat_history 
            WHERE username = ? AND thread_id = ? AND id < ?
            ORDER BY timestamp DESC 
            LIMIT ?
        """, (username, thread_id, message_id, context_size))
        
        before = [{'role': row[0], 'content': row[1], 'timestamp': row[2]} 
                 for row in reversed(cursor.fetchall())]
        
        # Nachrichten nach der gefundenen
        cursor.execute("""
            SELECT role, content, timestamp 
            FROM chat_history 
            WHERE username = ? AND thread_id = ? AND id > ?
            ORDER BY timestamp ASC 
            LIMIT ?
        """, (username, thread_id, message_id, context_size))
        
        after = [{'role': row[0], 'content': row[1], 'timestamp': row[2]} 
                for row in cursor.fetchall()]
        
        return {
            'before': before,
            'after': after
        }
        
    except Exception as e:
        logger.error(f"[SEARCH] Error getting context for message {message_id}: {e}")
        return {'before': [], 'after': []}
    finally:
        conn.close()

def highlight_search_term(text: str, search_term: str, max_length: int = 200) -> str:
    """
    Hebt Suchbegriff im Text hervor und kürzt bei Bedarf
    
    Args:
        text: Ursprungstext
        search_term: Zu suchender Begriff
        max_length: Max. Länge des Ergebnisses
    
    Returns:
        Text mit hervorgehobenen Suchbegriffen
    """
    if not search_term.strip():
        return text[:max_length] + "..." if len(text) > max_length else text
    
    # Case-insensitive highlight
    pattern = re.compile(re.escape(search_term), re.IGNORECASE)
    
    # Finde erste Übereinstimmung
    match = pattern.search(text)
    if not match:
        return text[:max_length] + "..." if len(text) > max_length else text
    
    # Extrahiere Snippet um die Übereinstimmung
    start = match.start()
    snippet_start = max(0, start - max_length // 2)
    snippet_end = min(len(text), snippet_start + max_length)
    
    snippet = text[snippet_start:snippet_end]
    
    # Hinzufügen von "..." wenn gekürzt
    if snippet_start > 0:
        snippet = "..." + snippet
    if snippet_end < len(text):
        snippet = snippet + "..."
    
    # Highlight hinzufügen
    highlighted = pattern.sub(
        r'<mark class="search-highlight">\g<0></mark>', 
        snippet
    )
    
    return highlighted

def get_search_suggestions(username: str, partial_query: str, limit: int = 10) -> List[str]:
    """
    Schlägt Suchbegriffe basierend auf häufigen Wörtern vor
    
    Args:
        username: Der User
        partial_query: Teilweise eingegebene Suche
        limit: Max. Anzahl Vorschläge
    
    Returns:
        Liste von Suchvorschlägen
    """
    if len(partial_query) < 2:
        return []
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Finde häufige Wörter die mit der Eingabe beginnen
        cursor.execute("""
            SELECT content FROM chat_history 
            WHERE username = ? AND content LIKE ?
            ORDER BY timestamp DESC 
            LIMIT 100
        """, (username, f"%{partial_query}%"))
        
        rows = cursor.fetchall()
        word_counts = {}
        
        # Extrahiere Wörter aus den Nachrichten
        for row in rows:
            content = row[0].lower()
            # Einfache Tokenisierung
            words = re.findall(r'\b\w+\b', content)
            
            for word in words:
                if (word.startswith(partial_query.lower()) and 
                    len(word) > len(partial_query) and
                    len(word) >= 3):  # Mindestlänge
                    word_counts[word] = word_counts.get(word, 0) + 1
        
        # Sortiere nach Häufigkeit
        suggestions = sorted(word_counts.items(), key=lambda x: x[1], reverse=True)
        return [word for word, count in suggestions[:limit]]
        
    except Exception as e:
        logger.error(f"[SEARCH] Error getting suggestions for {username}: {e}")
        return []
    finally:
        conn.close()

def get_search_statistics(username: str) -> Dict:
    """
    Holt Suchstatistiken für einen User
    
    Args:
        username: Der User
    
    Returns:
        Dict mit Statistiken
    """
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Gesamtanzahl Nachrichten
        cursor.execute("""
            SELECT COUNT(*) FROM chat_history WHERE username = ?
        """, (username,))
        total_messages = cursor.fetchone()[0]
        
        # Nachrichten nach Thread
        cursor.execute("""
            SELECT t.title, COUNT(h.id) as count
            FROM chat_history h
            LEFT JOIN chat_threads t ON h.thread_id = t.id AND h.username = t.username
            WHERE h.username = ?
            GROUP BY h.thread_id, t.title
            ORDER BY count DESC
        """, (username,))
        thread_stats = cursor.fetchall()
        
        # Nachrichten nach Datum (letzte 30 Tage)
        cursor.execute("""
            SELECT DATE(timestamp) as date, COUNT(*) as count
            FROM chat_history 
            WHERE username = ? AND timestamp >= datetime('now', '-30 days')
            GROUP BY DATE(timestamp)
            ORDER BY date DESC
        """, (username,))
        daily_stats = cursor.fetchall()
        
        # Rolle-Verteilung
        cursor.execute("""
            SELECT role, COUNT(*) as count
            FROM chat_history 
            WHERE username = ?
            GROUP BY role
        """, (username,))
        role_stats = cursor.fetchall()
        
        return {
            'total_messages': total_messages,
            'threads': [
                {'name': row[0] or 'Unbekannter Thread', 'count': row[1]} 
                for row in thread_stats
            ],
            'daily_activity': [
                {'date': row[0], 'count': row[1]} 
                for row in daily_stats
            ],
            'message_types': [
                {'role': row[0], 'count': row[1]} 
                for row in role_stats
            ]
        }
        
    except Exception as e:
        logger.error(f"[SEARCH] Error getting statistics for {username}: {e}")
        return {
            'total_messages': 0,
            'threads': [],
            'daily_activity': [],
            'message_types': []
        }
    finally:
        conn.close()

# ──────────────────────────────
# Advanced Search Functions
# ──────────────────────────────

def advanced_search(username: str, query: str, filters: Dict) -> List[Dict]:
    """
    Erweiterte Suche mit verschiedenen Filtern
    
    Args:
        username: Der User
        query: Suchbegriff
        filters: Dict mit Filtern (role, thread_id, date_from, date_to)
    
    Returns:
        Liste von Suchergebnissen
    """
    if not query.strip():
        return []
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        sql = """
            SELECT h.id, h.role, h.content, h.timestamp, h.thread_id,
                   t.title as thread_title
            FROM chat_history h
            LEFT JOIN chat_threads t ON h.thread_id = t.id AND h.username = t.username
            WHERE h.username = ?
        """
        params = [username]
        
        # Suchbegriff-Filter
        if '*' in query or '?' in query:
            # Wildcard-Suche
            search_pattern = query.replace('*', '%').replace('?', '_')
            sql += " AND h.content LIKE ?"
            params.append(search_pattern)
        elif '"' in query:
            # Exakte Phrase
            exact_phrase = query.strip('"')
            sql += " AND h.content LIKE ?"
            params.append(f"%{exact_phrase}%")
        else:
            # Standard-Suche (alle Wörter müssen vorkommen)
            words = query.split()
            for word in words:
                sql += " AND h.content LIKE ?"
                params.append(f"%{word}%")
        
        # Rollen-Filter
        if filters.get('role'):
            sql += " AND h.role = ?"
            params.append(filters['role'])
        
        # Thread-Filter
        if filters.get('thread_id'):
            sql += " AND h.thread_id = ?"
            params.append(filters['thread_id'])
        
        # Datumsbereich
        if filters.get('date_from'):
            sql += " AND h.timestamp >= ?"
            params.append(filters['date_from'])
        
        if filters.get('date_to'):
            sql += " AND h.timestamp <= ?"
            params.append(filters['date_to'])
        
        sql += " ORDER BY h.timestamp DESC LIMIT ?"
        params.append(filters.get('limit', 50))
        
        cursor.execute(sql, params)
        rows = cursor.fetchall()
        
        results = []
        for row in rows:
            context = get_message_context(username, row[0], row[4])
            
            results.append({
                'id': row[0],
                'role': row[1],
                'content': row[2],
                'timestamp': row[3],
                'thread_id': row[4],
                'thread_title': row[5] or 'Unbekannter Thread',
                'highlighted_content': highlight_search_term(row[2], query.strip('"')),
                'context': context
            })
        
        return results
        
    except Exception as e:
        logger.error(f"[SEARCH] Error in advanced search for {username}: {e}")
        return []
    finally:
        conn.close()

def export_search_results(username: str, results: List[Dict]) -> str:
    """
    Exportiert Suchergebnisse als CSV-String
    
    Args:
        username: Der User
        results: Suchergebnisse
    
    Returns:
        CSV-String
    """
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Header
    writer.writerow([
        'Timestamp', 'Thread', 'Role', 'Content', 'Thread_ID'
    ])
    
    # Daten
    for result in results:
        writer.writerow([
            result['timestamp'],
            result['thread_title'],
            result['role'],
            result['content'][:500],  # Begrenzte Länge
            result['thread_id']
        ])
    
    output.seek(0)
    return output.getvalue()

# ──────────────────────────────
# Cached Search Functions
# ──────────────────────────────

@cache_result("search_stats", ttl=600)
def get_search_statistics_cached(username: str):
    """Cached Version von get_search_statistics"""
    return get_search_statistics(username)

def invalidate_search_cache(username: str):
    """Such-bezogene Cache-Einträge löschen"""
    app_cache.delete(f"search_stats:{username}")

# ──────────────────────────────
# Search Helper Functions
# ──────────────────────────────

def parse_search_query(query: str) -> Dict:
    """
    Parst erweiterte Suchsyntax
    
    Args:
        query: Roh-Suchquery
    
    Returns:
        Dict mit geparseten Elementen
    """
    parsed = {
        'terms': [],
        'exact_phrases': [],
        'exclude_terms': [],
        'role_filter': None,
        'thread_filter': None
    }
    
    # Entferne spezielle Befehle
    parts = query.split()
    current_terms = []
    
    for part in parts:
        if part.startswith('role:'):
            parsed['role_filter'] = part[5:]
        elif part.startswith('thread:'):
            parsed['thread_filter'] = part[7:]
        elif part.startswith('-'):
            parsed['exclude_terms'].append(part[1:])
        elif '"' in part:
            # Sammle Phrase bis zum schließenden Anführungszeichen
            if part.count('"') == 2:
                parsed['exact_phrases'].append(part.strip('"'))
            else:
                current_terms.append(part)
        else:
            current_terms.append(part)
    
    parsed['terms'] = current_terms
    return parsed

def format_search_result_snippet(content: str, query: str, max_length: int = 150) -> str:
    """
    Formatiert Suchergebnis-Snippet mit Kontext
    
    Args:
        content: Vollständiger Inhalt
        query: Suchbegriff
        max_length: Maximale Snippet-Länge
    
    Returns:
        Formatiertes Snippet
    """
    if not query or len(content) <= max_length:
        return highlight_search_term(content, query, max_length)
    
    # Finde beste Position für Snippet
    query_lower = query.lower()
    content_lower = content.lower()
    
    pos = content_lower.find(query_lower)
    if pos == -1:
        return content[:max_length] + "..."
    
    # Zentriere um Suchbegriff
    start = max(0, pos - max_length // 2)
    end = min(len(content), start + max_length)
    
    snippet = content[start:end]
    
    if start > 0:
        snippet = "..." + snippet
    if end < len(content):
        snippet = snippet + "..."
    
    return highlight_search_term(snippet, query)

def set_user_subscription_tier_cached(username: str, tier: str):
    """set_user_subscription_tier mit Cache-Invalidierung"""
    set_user_subscription_tier(username, tier)
    invalidate_user_cache(username)
    app_cache.delete("all_users:")

def toggle_user_block_cached(username: str):
    """toggle_user_block mit Cache-Invalidierung"""
    toggle_user_block(username)
    invalidate_user_cache(username)
    app_cache.delete("all_users:")

def save_user_history(username: str, role: str, content: str):
    """Chat-Nachricht in DB speichern"""
    if not CHAT_TABLE_CREATED:
        return
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
        INSERT INTO chat_history (username, role, content) 
        VALUES (?, ?, ?)
    """, (username, role, content))
    conn.commit()
    conn.close()

def delete_user_history(username: str):
    """Chat-Verlauf löschen"""
    if not CHAT_TABLE_CREATED:
        return
        
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM chat_history WHERE username = ?", (username,))
    conn.commit()
    conn.close()

def toggle_user_block(username: str):
    """User blockieren/entblockieren"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT is_blocked FROM users WHERE username = ?", (username,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return
    current = row[0]
    new_status = 0 if current else 1
    cursor.execute("UPDATE users SET is_blocked = ? WHERE username = ?", 
                   (new_status, username))
    conn.commit()
    conn.close()

def delete_user_completely(username: str):
    """User und alle Daten löschen"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM users WHERE username = ?", (username,))
    cursor.execute("DELETE FROM chat_history WHERE username = ?", (username,))
    conn.commit()
    conn.close()

def get_user_persona(username: str) -> str:
    """Holt die gewählte Persona des Users"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("SELECT persona FROM users WHERE username = ?", (username,))
        result = cursor.fetchone()
        return result[0] if result and result[0] else "standard"
    except sqlite3.OperationalError:
        # Spalte existiert noch nicht
        return "standard"
    finally:
        conn.close()

def save_user_persona(username: str, persona: str):
    """Speichert die gewählte Persona des Users"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Prüfen ob users Tabelle persona Spalte hat, falls nicht hinzufügen
    cursor.execute("PRAGMA table_info(users)")
    columns = [column[1] for column in cursor.fetchall()]
    
    if 'persona' not in columns:
        cursor.execute("ALTER TABLE users ADD COLUMN persona TEXT DEFAULT 'standard'")
    
    cursor.execute("UPDATE users SET persona = ? WHERE username = ?", (persona, username))
    conn.commit()
    conn.close()

# ──────────────────────────────
# AI Model Definitions mit kreativen Namen
# ──────────────────────────────

class ModelProvider(Enum):
    OLLAMA = "ollama"
    OPENAI = "openai"

@dataclass
class AIModel:
    """Definition eines KI-Modells mit kreativem Namen"""
    id: str                    # Technische ID
    display_name: str          # Kreativer Anzeigename
    provider: ModelProvider    # Anbieter (Ollama/OpenAI)
    model_name: str           # Echter Model-Name bei API
    description: str          # Beschreibung
    max_tokens: int           # Token-Limit
    cost_per_1k_tokens: float # Kosten (falls relevant)
    capabilities: List[str]    # Fähigkeiten
    subscription_tier: str     # Mindest-Tier
    is_active: bool           # Verfügbar?
    personality: str          # Persönlichkeitsbeschreibung

# Verfügbare KI-Modelle mit kreativen Namen
AI_MODELS = {
    # Ollama-basierte Modelle (lokal/kostenlos)
    "nexus": AIModel(
        id="nexus",
        display_name="Nexus",
        provider=ModelProvider.OLLAMA,
        model_name="llama2:7b",
        description="Ein zuverlässiger Allround-Assistent mit ausgeglichener Persönlichkeit",
        max_tokens=4096,
        cost_per_1k_tokens=0.0,
        capabilities=["general_chat", "coding", "analysis"],
        subscription_tier="free",
        is_active=True,
        personality="Ausgewogen, hilfsbereit, logisch denkend"
    ),
    
    "aurora": AIModel(
        id="aurora",
        display_name="Aurora",
        provider=ModelProvider.OLLAMA,
        model_name="llama2:13b",
        description="Eine kreative und intuitive KI für komplexe Aufgaben",
        max_tokens=4096,
        cost_per_1k_tokens=0.0,
        capabilities=["creative_writing", "brainstorming", "complex_reasoning"],
        subscription_tier="pro",
        is_active=True,
        personality="Kreativ, einfühlsam, visionär"
    ),
    
    "cipher": AIModel(
        id="cipher",
        display_name="Cipher",
        provider=ModelProvider.OLLAMA,
        model_name="codellama:13b",
        description="Spezialist für Programmierung und technische Analysen",
        max_tokens=4096,
        cost_per_1k_tokens=0.0,
        capabilities=["coding", "debugging", "technical_analysis", "architecture"],
        subscription_tier="pro",
        is_active=True,
        personality="Präzise, methodisch, technisch versiert"
    ),
    
    # OpenAI-basierte Modelle (kostenpflichtig)
    "phoenix": AIModel(
        id="phoenix",
        display_name="Phoenix",
        provider=ModelProvider.OPENAI,
        model_name="gpt-3.5-turbo",
        description="Schneller und effizienter Conversational AI",
        max_tokens=4096,
        cost_per_1k_tokens=0.002,
        capabilities=["general_chat", "quick_responses", "multi_language"],
        subscription_tier="pro",
        is_active=True,
        personality="Schnell, direkt, vielseitig"
    ),
    
    "prometheus": AIModel(
        id="prometheus",
        display_name="Prometheus",
        provider=ModelProvider.OPENAI,
        model_name="gpt-4",
        description="Hochintelligente KI für komplexeste Aufgaben und tiefe Analysen",
        max_tokens=8192,
        cost_per_1k_tokens=0.03,
        capabilities=["advanced_reasoning", "research", "complex_analysis", "creative_writing"],
        subscription_tier="premium",
        is_active=True,
        personality="Tiefgreifend, analytisch, weise"
    ),
    
    "atlas": AIModel(
        id="atlas",
        display_name="Atlas",
        provider=ModelProvider.OPENAI,
        model_name="gpt-4-turbo",
        description="Kraftvolle KI mit erweiterten Fähigkeiten für anspruchsvolle Tasks",
        max_tokens=128000,
        cost_per_1k_tokens=0.01,
        capabilities=["long_context", "document_analysis", "research", "multi_modal"],
        subscription_tier="premium",
        is_active=True,
        personality="Kraftvoll, gründlich, umfassend"
    )
}

# ──────────────────────────────
# Abstract AI Provider Interface
# ──────────────────────────────

class AIProvider(ABC):
    """Abstrakte Basis-Klasse für KI-Anbieter"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.is_available = False
        
    @abstractmethod
    async def test_connection(self) -> bool:
        """Testet Verbindung zum Anbieter"""
        pass
    
    @abstractmethod
    async def generate_response(self, messages: List[Dict], model: AIModel, **kwargs) -> Dict:
        """Generiert Antwort vom Model"""
        pass
    
    @abstractmethod
    def get_available_models(self) -> List[str]:
        """Listet verfügbare Modelle auf"""
        pass

# ──────────────────────────────
# Ollama Provider Implementation
# ──────────────────────────────

class OllamaProvider(AIProvider):
    """Ollama-Anbieter für lokale Modelle"""
    
    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.base_url = config.get('base_url', 'http://localhost:11434')
        self.timeout = config.get('timeout', 120)
        
    async def test_connection(self) -> bool:
        """Testet Ollama-Verbindung"""
        try:
            async with httpx.AsyncClient(timeout=self.timeout) as client:
                response = await client.get(f"{self.base_url}/api/tags")
                self.is_available = response.status_code == 200
                return self.is_available
        except Exception as e:
            logger.error(f"[OLLAMA] Connection test failed: {e}")
            self.is_available = False
            return False
    
    async def generate_response(self, messages: List[Dict], model: AIModel, **kwargs) -> Dict:
        """Generiert Antwort über Ollama"""
        start_time = time.time()
        
        try:
            # Format für Ollama API
            prompt = self._format_messages_for_ollama(messages)
            
            async with httpx.AsyncClient(timeout=self.timeout) as client:
                payload = {
                    "model": model.model_name,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": kwargs.get('temperature', 0.7),
                        "top_p": kwargs.get('top_p', 0.9),
                        "max_tokens": kwargs.get('max_tokens', model.max_tokens)
                    }
                }
                
                response = await client.post(
                    f"{self.base_url}/api/generate",
                    json=payload
                )
                
                if response.status_code == 200:
                    data = response.json()
                    duration = time.time() - start_time
                    
                    return {
                        'success': True,
                        'content': data.get('response', ''),
                        'model_used': model.display_name,
                        'provider': 'ollama',
                        'duration': duration,
                        'tokens_used': len(data.get('response', '').split()) * 1.3,  # Schätzung
                        'cost': 0.0,  # Ollama ist kostenlos
                        'metadata': {
                            'model_name': model.model_name,
                            'done': data.get('done', False)
                        }
                    }
                else:
                    return {
                        'success': False,
                        'error': f"Ollama API error: {response.status_code}",
                        'model_used': model.display_name
                    }
                    
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'model_used': model.display_name
            }
    
    def _format_messages_for_ollama(self, messages: List[Dict]) -> str:
        """Konvertiert Chat-Messages zu Ollama-Prompt"""
        formatted_parts = []
        
        for message in messages:
            role = message.get('role', 'user')
            content = message.get('content', '')
            
            if role == 'system':
                formatted_parts.append(f"System: {content}")
            elif role == 'user':
                formatted_parts.append(f"Human: {content}")
            elif role == 'assistant':
                formatted_parts.append(f"Assistant: {content}")
        
        # Abschluss für Antwort
        formatted_parts.append("Assistant:")
        return "\n\n".join(formatted_parts)
    
    def get_available_models(self) -> List[str]:
        """Holt verfügbare Ollama-Modelle"""
        try:
            import requests
            response = requests.get(f"{self.base_url}/api/tags", timeout=10)
            if response.status_code == 200:
                data = response.json()
                return [model['name'] for model in data.get('models', [])]
        except Exception as e:
            logger.error(f"[OLLAMA] Error fetching models: {e}")
        return []

# ──────────────────────────────
# OpenAI Provider Implementation
# ──────────────────────────────

class OpenAIProvider(AIProvider):
    """OpenAI-Anbieter für GPT-Modelle"""
    
    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.api_key = config.get('api_key', os.getenv('OPENAI_API_KEY'))
        self.organization = config.get('organization')
        
        # OpenAI Client konfigurieren
        openai.api_key = self.api_key
        if self.organization:
            openai.organization = self.organization
    
    async def test_connection(self) -> bool:
        """Testet OpenAI-Verbindung"""
        try:
            # Einfacher Test-Request
            response = await openai.ChatCompletion.acreate(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": "Test"}],
                max_tokens=5
            )
            self.is_available = True
            return True
        except Exception as e:
            logger.error(f"[OPENAI] Connection test failed: {e}")
            self.is_available = False
            return False
    
    async def generate_response(self, messages: List[Dict], model: AIModel, **kwargs) -> Dict:
        """Generiert Antwort über OpenAI API"""
        start_time = time.time()
        
        try:
            response = await openai.ChatCompletion.acreate(
                model=model.model_name,
                messages=messages,
                max_tokens=kwargs.get('max_tokens', min(model.max_tokens, 2000)),
                temperature=kwargs.get('temperature', 0.7),
                top_p=kwargs.get('top_p', 0.9),
                presence_penalty=kwargs.get('presence_penalty', 0),
                frequency_penalty=kwargs.get('frequency_penalty', 0)
            )
            
            duration = time.time() - start_time
            usage = response.usage
            
            # Kosten berechnen
            tokens_used = usage.total_tokens
            cost = (tokens_used / 1000) * model.cost_per_1k_tokens
            
            return {
                'success': True,
                'content': response.choices[0].message.content,
                'model_used': model.display_name,
                'provider': 'openai',
                'duration': duration,
                'tokens_used': tokens_used,
                'cost': cost,
                'metadata': {
                    'model_name': model.model_name,
                    'prompt_tokens': usage.prompt_tokens,
                    'completion_tokens': usage.completion_tokens,
                    'finish_reason': response.choices[0].finish_reason
                }
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'model_used': model.display_name
            }
    
    def get_available_models(self) -> List[str]:
        """Holt verfügbare OpenAI-Modelle"""
        try:
            models = openai.Model.list()
            return [model.id for model in models.data if 'gpt' in model.id]
        except Exception as e:
            logger.error(f"[OPENAI] Error fetching models: {e}")
        return []

# ──────────────────────────────
# Model Manager
# ──────────────────────────────

class ModelManager:
    """Zentraler Manager für alle KI-Modelle"""
    
    def __init__(self):
        self.providers: Dict[ModelProvider, AIProvider] = {}
        self.models = AI_MODELS
        self.default_model = "nexus"  # Fallback-Model
        
        # Provider initialisieren
        self._initialize_providers()
    
    def _initialize_providers(self):
        """Initialisiert alle verfügbaren Provider"""
        # Ollama Provider
        ollama_config = {
            'base_url': os.getenv('OLLAMA_BASE_URL', 'http://localhost:11434'),
            'timeout': 120
        }
        self.providers[ModelProvider.OLLAMA] = OllamaProvider(ollama_config)
        
        # OpenAI Provider
        openai_config = {
            'api_key': os.getenv('OPENAI_API_KEY'),
            'organization': os.getenv('OPENAI_ORGANIZATION')
        }
        self.providers[ModelProvider.OPENAI] = OpenAIProvider(openai_config)
    
    async def test_all_providers(self) -> Dict[str, bool]:
        """Testet alle Provider"""
        results = {}
        for provider_name, provider in self.providers.items():
            results[provider_name.value] = await provider.test_connection()
        return results
    
    def get_available_models_for_user(self, username: str) -> Dict[str, AIModel]:
        """Holt verfügbare Modelle basierend auf User-Subscription"""
        user_tier = get_user_subscription_tier_cached(username)
        
        available_models = {}
        for model_id, model in self.models.items():
            if not model.is_active:
                continue
                
            # Subscription-Check
            tier_levels = {"free": 1, "pro": 2, "premium": 3}
            user_level = tier_levels.get(user_tier, 1)
            model_level = tier_levels.get(model.subscription_tier, 1)
            
            if user_level >= model_level:
                # Provider-Verfügbarkeit prüfen
                provider = self.providers.get(model.provider)
                if provider and provider.is_available:
                    available_models[model_id] = model
        
        return available_models
    
    async def generate_response(self, model_id: str, messages: List[Dict], 
                              username: str, **kwargs) -> Dict:
        """Generiert Antwort mit spezifischem Model"""
        # Model validieren
        if model_id not in self.models:
            model_id = self.default_model
        
        model = self.models[model_id]
        
        # User-Berechtigung prüfen
        available_models = self.get_available_models_for_user(username)
        if model_id not in available_models:
            model_id = self.default_model
            model = self.models[model_id]
        
        # Provider holen
        provider = self.providers.get(model.provider)
        if not provider or not provider.is_available:
            # Fallback zu anderem Provider
            fallback_result = await self._try_fallback_models(messages, username, **kwargs)
            if fallback_result:
                return fallback_result
        
        # Response generieren
        try:
            result = await provider.generate_response(messages, model, **kwargs)
            
            # Usage tracking
            if result.get('success'):
                await self._track_model_usage(username, model_id, result)
            
            return result
            
        except Exception as e:
            logger.error(f"[MODEL] Error with {model.display_name}: {e}")
            return await self._try_fallback_models(messages, username, **kwargs)
    
    async def _try_fallback_models(self, messages: List[Dict], username: str, **kwargs) -> Dict:
        """Versucht Fallback-Modelle bei Fehlern"""
        available_models = self.get_available_models_for_user(username)
        
        # Prioritätsliste für Fallbacks
        fallback_priority = ["nexus", "aurora", "phoenix", "cipher"]
        
        for model_id in fallback_priority:
            if model_id in available_models:
                try:
                    model = self.models[model_id]
                    provider = self.providers.get(model.provider)
                    
                    if provider and provider.is_available:
                        result = await provider.generate_response(messages, model, **kwargs)
                        if result.get('success'):
                            result['is_fallback'] = True
                            return result
                except Exception:
                    continue
        
        # Wenn alle Fallbacks fehlschlagen
        return {
            'success': False,
            'error': 'Alle KI-Modelle sind momentan nicht verfügbar',
            'model_used': 'none'
        }
    
    async def _track_model_usage(self, username: str, model_id: str, result: Dict):
        """Trackt Model-Usage für Statistiken"""
        try:
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            
            # Usage-Tabelle erstellen falls nicht vorhanden
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS model_usage (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT NOT NULL,
                    model_id TEXT NOT NULL,
                    model_name TEXT NOT NULL,
                    provider TEXT NOT NULL,
                    tokens_used INTEGER DEFAULT 0,
                    cost REAL DEFAULT 0.0,
                    duration REAL DEFAULT 0.0,
                    success INTEGER DEFAULT 1,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (username) REFERENCES users (username)
                )
            """)
            
            # Usage-Eintrag hinzufügen
            cursor.execute("""
                INSERT INTO model_usage 
                (username, model_id, model_name, provider, tokens_used, cost, duration, success)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                username,
                model_id,
                result.get('model_used', 'unknown'),
                result.get('provider', 'unknown'),
                result.get('tokens_used', 0),
                result.get('cost', 0.0),
                result.get('duration', 0.0),
                1 if result.get('success') else 0
            ))
            
            conn.commit()
            conn.close()
            
        except Exception as e:
            logger.error(f"[MODEL] Usage tracking error: {e}")

# Globaler Model-Manager
model_manager = ModelManager()

# ──────────────────────────────
# Enhanced Subscription Functions
# ──────────────────────────────
def get_user_subscription_tier(username: str) -> str:
    """Holt das Subscription-Tier des Users"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("SELECT subscription_tier FROM users WHERE username = ?", (username,))
        result = cursor.fetchone()
        return result[0] if result and result[0] else "free"
    except sqlite3.OperationalError:
        # Spalte existiert noch nicht - hinzufügen
        cursor.execute("ALTER TABLE users ADD COLUMN subscription_tier TEXT DEFAULT 'free'")
        conn.commit()
        return "free"
    finally:
        conn.close()

def set_user_subscription_tier(username: str, tier: str):
    """Setzt das Subscription-Tier für einen User"""
    if tier not in SUBSCRIPTION_TIERS:
        tier = "free"
    
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    # Sicherstellen, dass Spalte existiert
    try:
        cursor.execute("UPDATE users SET subscription_tier = ? WHERE username = ?", (tier, username))
    except sqlite3.OperationalError:
        cursor.execute("ALTER TABLE users ADD COLUMN subscription_tier TEXT DEFAULT 'free'")
        cursor.execute("UPDATE users SET subscription_tier = ? WHERE username = ?", (tier, username))
    
    conn.commit()
    conn.close()

def get_available_personas_for_user(username: str) -> dict:
    """Gibt verfügbare Personas basierend auf Subscription zurück"""
    user_tier = get_user_subscription_tier(username)
    tier_config = SUBSCRIPTION_TIERS.get(user_tier, SUBSCRIPTION_TIERS["free"])
    available_persona_keys = tier_config["available_personas"]
    
    return {key: persona for key, persona in PERSONAS.items() 
            if key in available_persona_keys}

def check_enhanced_rate_limit(username: str) -> dict:
    """Erweiterte Rate-Limit-Prüfung basierend auf Subscription"""
    user_tier = get_user_subscription_tier(username)
    tier_config = SUBSCRIPTION_TIERS.get(user_tier, SUBSCRIPTION_TIERS["free"])
    
    max_per_hour = tier_config["max_messages_per_hour"]
    max_per_day = tier_config["max_messages_per_day"]
    
    # Unlimited für Premium
    if max_per_hour == -1:
        return {"allowed": True, "remaining_hour": "∞", "remaining_day": "∞"}
    
    limits = load_rate_limits()
    current_time = _pytime.time()
    
    user_data = limits.get(username, {"messages": [], "daily_messages": [], "last_reset": current_time})
    
    # Alte Nachrichten entfernen
    hour_ago = current_time - 3600
    day_ago = current_time - 86400
    
    user_data["messages"] = [msg_time for msg_time in user_data["messages"] if msg_time > hour_ago]
    user_data["daily_messages"] = [msg_time for msg_time in user_data.get("daily_messages", []) if msg_time > day_ago]
    
    # Prüfen
    hourly_used = len(user_data["messages"])
    daily_used = len(user_data["daily_messages"])
    
    if hourly_used >= max_per_hour or (max_per_day != -1 and daily_used >= max_per_day):
        return {
            "allowed": False, 
            "remaining_hour": max_per_hour - hourly_used,
            "remaining_day": max_per_day - daily_used if max_per_day != -1 else "∞",
            "tier": user_tier
        }
    
    # Neue Nachricht hinzufügen
    user_data["messages"].append(current_time)
    user_data["daily_messages"].append(current_time)
    limits[username] = user_data
    save_rate_limits(limits)
    
    return {
        "allowed": True,
        "remaining_hour": max_per_hour - hourly_used - 1,
        "remaining_day": max_per_day - daily_used - 1 if max_per_day != -1 else "∞",
        "tier": user_tier
    }

def get_response_with_context(current_message: str, chat_history: list, persona: str = "standard") -> str:
    """
    Holt KI-Antwort mit Chat-Kontext und Persona
    """
    # Chat-Historie in das richtige Format für Ollama konvertieren
    messages = []
    
    # System-Prompt basierend auf gewählter Persona
    persona_config = PERSONAS.get(persona, PERSONAS["standard"])
    messages.append({
        "role": "system", 
        "content": persona_config["system_prompt"]
    })
    
    # Chat-Historie hinzufügen (letzte 20 Nachrichten für besseren Kontext)
    for msg in chat_history[-20:]:
        if msg["role"] in ["user", "assistant"]:
            messages.append({
                "role": msg["role"],
                "content": msg["content"]
            })
    
    # Aktuelle Nachricht hinzufügen
    messages.append({
        "role": "user",
        "content": current_message
    })
    
    # An get_response weitergeben
    return get_response_with_messages(messages)

# ──────────────────────────────
# File Types and Configuration
# ──────────────────────────────

class FileType(Enum):
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"

ALLOWED_MIME_TYPES = {
    'application/pdf': FileType.PDF,
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': FileType.DOCX,
    'text/plain': FileType.TXT,
    'image/jpeg': FileType.IMAGE,
    'image/png': FileType.IMAGE,
    'image/gif': FileType.IMAGE,
    'image/bmp': FileType.IMAGE,
    'image/tiff': FileType.IMAGE
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
MAX_FILES_PER_USER = 100
UPLOAD_DIR = "uploads"
CHUNK_SIZE = 1000  # Zeichen pro Text-Chunk

@dataclass
class ProcessedFile:
    """Information über eine verarbeitete Datei"""
    id: str
    original_name: str
    file_type: FileType
    file_size: int
    upload_date: datetime
    text_content: str
    chunks: List[str]
    metadata: Dict
    username: str
    thread_id: Optional[str] = None

@dataclass
class FileChunk:
    """Ein Text-Chunk aus einer Datei"""
    chunk_id: str
    file_id: str
    content: str
    chunk_index: int
    char_count: int

# ──────────────────────────────
# File Storage Management
# ──────────────────────────────

def ensure_upload_directory():
    """Stellt sicher dass Upload-Verzeichnis existiert"""
    upload_path = Path(UPLOAD_DIR)
    upload_path.mkdir(exist_ok=True)
    return upload_path

def generate_file_id(username: str, filename: str) -> str:
    """Generiert eindeutige File-ID"""
    timestamp = datetime.now().isoformat()
    content = f"{username}:{filename}:{timestamp}"
    return hashlib.md5(content.encode()).hexdigest()[:16]

def get_file_path(file_id: str) -> Path:
    """Holt Dateipfad für File-ID"""
    upload_path = ensure_upload_directory()
    return upload_path / f"{file_id}.dat"

def detect_file_type(file_content: bytes, filename: str) -> Tuple[FileType, str]:
    """Erkennt Dateityp basierend auf Inhalt und Name"""
    try:
        # Magic number detection
        mime_type = magic.from_buffer(file_content, mime=True)
        
        if mime_type in ALLOWED_MIME_TYPES:
            return ALLOWED_MIME_TYPES[mime_type], mime_type
        
        # Fallback auf Dateiendung
        _, ext = os.path.splitext(filename.lower())
        if ext == '.pdf':
            return FileType.PDF, 'application/pdf'
        elif ext == '.docx':
            return FileType.DOCX, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif ext == '.txt':
            return FileType.TXT, 'text/plain'
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            return FileType.IMAGE, f'image/{ext[1:]}'
        
    except Exception as e:
        logger.error(f"[FILES] File type detection error: {e}")
    
    return FileType.UNKNOWN, 'application/octet-stream'

# ──────────────────────────────
# Text Extraction
# ──────────────────────────────

def extract_text_from_pdf(file_path: Path) -> Tuple[str, Dict]:
    """Extrahiert Text aus PDF-Datei"""
    try:
        text_content = ""
        metadata = {"pages": 0, "extraction_method": "PyPDF2"}
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata["pages"] = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Seite {page_num + 1} ---\n"
                        text_content += page_text
                        text_content += "\n"
                except Exception as e:
                    logger.warning(f"[FILES] Error extracting page {page_num}: {e}")
                    continue
        
        # PDF-Metadaten
        try:
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get('/Title', ''),
                    "author": pdf_reader.metadata.get('/Author', ''),
                    "subject": pdf_reader.metadata.get('/Subject', ''),
                    "creator": pdf_reader.metadata.get('/Creator', '')
                })
        except:
            pass
        
        return text_content.strip(), metadata
        
    except Exception as e:
        logger.error(f"[FILES] PDF extraction error: {e}")
        return "", {"error": str(e), "extraction_method": "failed"}

def extract_text_from_docx(file_path: Path) -> Tuple[str, Dict]:
    """Extrahiert Text aus DOCX-Datei"""
    try:
        doc = Document(file_path)
        
        text_content = ""
        metadata = {
            "paragraphs": 0,
            "tables": 0,
            "extraction_method": "python-docx"
        }
        
        # Paragraphen extrahieren
        for para in doc.paragraphs:
            if para.text.strip():
                text_content += para.text + "\n"
                metadata["paragraphs"] += 1
        
        # Tabellen extrahieren
        for table in doc.tables:
            metadata["tables"] += 1
            text_content += "\n--- Tabelle ---\n"
            
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text_content += " | ".join(row_text) + "\n"
            text_content += "\n"
        
        # Core Properties
        try:
            core_props = doc.core_properties
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            })
        except:
            pass
        
        return text_content.strip(), metadata
        
    except Exception as e:
        logger.error(f"[FILES] DOCX extraction error: {e}")
        return "", {"error": str(e), "extraction_method": "failed"}

def extract_text_from_image(file_path: Path) -> Tuple[str, Dict]:
    """Extrahiert Text aus Bild mit OCR"""
    try:
        # OCR mit Tesseract
        image = Image.open(file_path)
        text_content = pytesseract.image_to_string(image, lang='deu+eng')
        
        metadata = {
            "extraction_method": "tesseract_ocr",
            "image_size": image.size,
            "image_mode": image.mode,
            "format": image.format
        }
        
        return text_content.strip(), metadata
        
    except Exception as e:
        logger.error(f"[FILES] Image OCR error: {e}")
        return "", {"error": str(e), "extraction_method": "ocr_failed"}

def extract_text_from_txt(file_path: Path) -> Tuple[str, Dict]:
    """Extrahiert Text aus TXT-Datei"""
    try:
        # Verschiedene Encodings versuchen
        encodings = ['utf-8', 'utf-16', 'latin1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text_content = file.read()
                
                metadata = {
                    "extraction_method": "text_file",
                    "encoding": encoding,
                    "line_count": text_content.count('\n') + 1
                }
                
                return text_content.strip(), metadata
                
            except UnicodeDecodeError:
                continue
        
        # Wenn alle Encodings fehlschlagen
        return "", {"error": "Encoding detection failed", "extraction_method": "failed"}
        
    except Exception as e:
        logger.error(f"[FILES] TXT extraction error: {e}")
        return "", {"error": str(e), "extraction_method": "failed"}

# ──────────────────────────────
# Text Chunking
# ──────────────────────────────

def create_text_chunks(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = 100) -> List[str]:
    """Teilt Text in überlappende Chunks für besseren Kontext"""
    if not text or len(text) <= chunk_size:
        return [text] if text else []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Bei letztem Chunk: nehme alles was übrig ist
        if end >= len(text):
            chunk = text[start:]
            if chunk.strip():
                chunks.append(chunk.strip())
            break
        
        # Versuche bei Wort-/Satzgrenze zu trennen
        chunk_text = text[start:end]
        
        # Suche nach bestem Trennpunkt
        best_split = end
        for delimiter in ['\n\n', '. ', '\n', ', ']:
            last_occurrence = chunk_text.rfind(delimiter)
            if last_occurrence > chunk_size * 0.7:  # Mindestens 70% des Chunks
                best_split = start + last_occurrence + len(delimiter)
                break
        
        chunk = text[start:best_split].strip()
        if chunk:
            chunks.append(chunk)
        
        # Nächster Start mit Überlappung
        start = best_split - overlap
        if start < 0:
            start = best_split
    
    return chunks

def create_smart_summary(text: str, max_length: int = 500) -> str:
    """Erstellt intelligente Zusammenfassung eines Textes"""
    if len(text) <= max_length:
        return text
    
    # Erste Sätze nehmen bis max_length erreicht
    sentences = text.split('. ')
    summary_parts = []
    current_length = 0
    
    for sentence in sentences:
        sentence_with_period = sentence + '. ' if not sentence.endswith('.') else sentence + ' '
        
        if current_length + len(sentence_with_period) <= max_length:
            summary_parts.append(sentence_with_period)
            current_length += len(sentence_with_period)
        else:
            break
    
    summary = ''.join(summary_parts).strip()
    
    if current_length < len(text):
        summary += "..."
    
    return summary

# ──────────────────────────────
# File Processing Pipeline
# ──────────────────────────────

def process_uploaded_file(file_content: bytes, filename: str, username: str, 
                         thread_id: Optional[str] = None) -> ProcessedFile:
    """Hauptfunktion zur Dateiverarbeitung"""
    
    # File-ID generieren
    file_id = generate_file_id(username, filename)
    
    # Dateityp erkennen
    file_type, mime_type = detect_file_type(file_content, filename)
    
    if file_type == FileType.UNKNOWN:
        raise ValueError(f"Dateityp nicht unterstützt: {mime_type}")
    
    # Datei temporär speichern
    file_path = get_file_path(file_id)
    
    try:
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        # Text extrahieren
        text_content = ""
        metadata = {"mime_type": mime_type}
        
        if file_type == FileType.PDF:
            text_content, extraction_metadata = extract_text_from_pdf(file_path)
        elif file_type == FileType.DOCX:
            text_content, extraction_metadata = extract_text_from_docx(file_path)
        elif file_type == FileType.TXT:
            text_content, extraction_metadata = extract_text_from_txt(file_path)
        elif file_type == FileType.IMAGE:
            text_content, extraction_metadata = extract_text_from_image(file_path)
        else:
            raise ValueError(f"Unbekannter Dateityp: {file_type}")
        
        metadata.update(extraction_metadata)
        
        # Text-Chunks erstellen
        chunks = create_text_chunks(text_content) if text_content else []
        
        # Summary erstellen
        summary = create_smart_summary(text_content) if text_content else "Keine Textinhalte gefunden"
        metadata["summary"] = summary
        metadata["chunks_count"] = len(chunks)
        metadata["text_length"] = len(text_content)
        
        # ProcessedFile-Objekt erstellen
        processed_file = ProcessedFile(
            id=file_id,
            original_name=filename,
            file_type=file_type,
            file_size=len(file_content),
            upload_date=datetime.now(),
            text_content=text_content,
            chunks=chunks,
            metadata=metadata,
            username=username,
            thread_id=thread_id
        )
        
        return processed_file
        
    except Exception as e:
        # Aufräumen bei Fehler
        if file_path.exists():
            file_path.unlink()
        raise e

def validate_file_upload(file_content: bytes, filename: str, username: str) -> Tuple[bool, str]:
    """Validiert Upload-Datei"""
    
    # Größe prüfen
    if len(file_content) > MAX_FILE_SIZE:
        return False, f"Datei zu groß (max. {MAX_FILE_SIZE // (1024*1024)} MB)"
    
    if len(file_content) == 0:
        return False, "Datei ist leer"
    
    # Dateityp prüfen
    file_type, mime_type = detect_file_type(file_content, filename)
    if file_type == FileType.UNKNOWN:
        return False, f"Dateityp nicht unterstützt: {mime_type}"
    
    # User-Limit prüfen
    user_file_count = get_user_file_count(username)
    if user_file_count >= MAX_FILES_PER_USER:
        return False, f"Maximale Anzahl Dateien erreicht ({MAX_FILES_PER_USER})"
    
    return True, "OK"

def get_user_file_count(username: str) -> int:
    """Zählt Dateien eines Users"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        cursor.execute("SELECT COUNT(*) FROM uploaded_files WHERE username = ?", (username,))
        count = cursor.fetchone()[0]
        
        conn.close()
        return count
        
    except sqlite3.OperationalError:
        return 0  # Tabelle existiert noch nicht

# mach hier das aus dem zweiten artefakt fpr databesfunktion oder wie auch immer 
# ──────────────────────────────
# Database Integration für File Processing
# ──────────────────────────────

def upgrade_database_for_files():
    """Erweitert die Datenbank um File-Processing-Support"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Uploaded Files Tabelle
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS uploaded_files (
                id TEXT PRIMARY KEY,
                username TEXT NOT NULL,
                thread_id TEXT,
                original_name TEXT NOT NULL,
                file_type TEXT NOT NULL,
                file_size INTEGER NOT NULL,
                mime_type TEXT,
                upload_date DATETIME DEFAULT CURRENT_TIMESTAMP,
                text_content TEXT,
                summary TEXT,
                metadata TEXT,
                chunks_count INTEGER DEFAULT 0,
                is_processed INTEGER DEFAULT 1,
                FOREIGN KEY (username) REFERENCES users (username),
                FOREIGN KEY (thread_id) REFERENCES chat_threads (id)
            )
        """)
        
        # File Chunks Tabelle für bessere Suche
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS file_chunks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_id TEXT NOT NULL,
                chunk_index INTEGER NOT NULL,
                content TEXT NOT NULL,
                char_count INTEGER DEFAULT 0,
                FOREIGN KEY (file_id) REFERENCES uploaded_files (id)
            )
        """)
        
        # File Usage Tracking
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS file_usage (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                file_id TEXT NOT NULL,
                username TEXT NOT NULL,
                action TEXT NOT NULL,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                context TEXT,
                FOREIGN KEY (file_id) REFERENCES uploaded_files (id),
                FOREIGN KEY (username) REFERENCES users (username)
            )
        """)
        
        conn.commit()
        logger.info("[FILES] Database upgraded for file processing")
        
    except Exception as e:
        logger.error(f"[FILES] Database upgrade error: {e}")
        conn.rollback()
        raise
    finally:
        conn.close()

# ──────────────────────────────
# File Database Operations
# ──────────────────────────────

def save_processed_file_to_db(processed_file: ProcessedFile) -> bool:
    """Speichert verarbeitete Datei in der Datenbank"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Hauptdatei-Eintrag
        cursor.execute("""
            INSERT INTO uploaded_files 
            (id, username, thread_id, original_name, file_type, file_size, 
             mime_type, text_content, summary, metadata, chunks_count)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            processed_file.id,
            processed_file.username,
            processed_file.thread_id,
            processed_file.original_name,
            processed_file.file_type.value,
            processed_file.file_size,
            processed_file.metadata.get('mime_type', ''),
            processed_file.text_content,
            processed_file.metadata.get('summary', ''),
            json.dumps(processed_file.metadata),
            len(processed_file.chunks)
        ))
        
        # Text-Chunks speichern
        for i, chunk in enumerate(processed_file.chunks):
            cursor.execute("""
                INSERT INTO file_chunks (file_id, chunk_index, content, char_count)
                VALUES (?, ?, ?, ?)
            """, (processed_file.id, i, chunk, len(chunk)))
        
        # Usage-Event loggen
        log_file_usage(processed_file.id, processed_file.username, "upload", cursor)
        
        conn.commit()
        logger.info(f"[FILES] File saved to database: {processed_file.id}")
        return True
        
    except Exception as e:
        logger.error(f"[FILES] Database save error: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def get_user_files(username: str, thread_id: Optional[str] = None) -> List[Dict]:
    """Holt Dateien eines Users"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        if thread_id:
            cursor.execute("""
                SELECT id, original_name, file_type, file_size, upload_date, 
                       summary, chunks_count, metadata
                FROM uploaded_files 
                WHERE username = ? AND thread_id = ?
                ORDER BY upload_date DESC
            """, (username, thread_id))
        else:
            cursor.execute("""
                SELECT id, original_name, file_type, file_size, upload_date, 
                       summary, chunks_count, metadata
                FROM uploaded_files 
                WHERE username = ?
                ORDER BY upload_date DESC
            """, (username,))
        
        rows = cursor.fetchall()
        
        files = []
        for row in rows:
            try:
                metadata = json.loads(row[7]) if row[7] else {}
            except:
                metadata = {}
            
            files.append({
                'id': row[0],
                'original_name': row[1],
                'file_type': row[2],
                'file_size': row[3],
                'upload_date': row[4],
                'summary': row[5],
                'chunks_count': row[6],
                'metadata': metadata
            })
        
        return files
        
    except sqlite3.OperationalError:
        return []
    finally:
        conn.close()

def get_file_by_id(file_id: str, username: str) -> Optional[ProcessedFile]:
    """Holt Datei-Details mit Berechtigung-Check"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("""
            SELECT id, username, thread_id, original_name, file_type, file_size,
                   upload_date, text_content, summary, metadata, chunks_count
            FROM uploaded_files 
            WHERE id = ? AND username = ?
        """, (file_id, username))
        
        row = cursor.fetchone()
        if not row:
            return None
        
        # Chunks laden
        cursor.execute("""
            SELECT content FROM file_chunks 
            WHERE file_id = ? 
            ORDER BY chunk_index
        """, (file_id,))
        
        chunk_rows = cursor.fetchall()
        chunks = [chunk_row[0] for chunk_row in chunk_rows]
        
        # Metadata parsen
        try:
            metadata = json.loads(row[9]) if row[9] else {}
        except:
            metadata = {}
        
        # ProcessedFile rekonstruieren
        processed_file = ProcessedFile(
            id=row[0],
            username=row[1],
            thread_id=row[2],
            original_name=row[3],
            file_type=FileType(row[4]),
            file_size=row[5],
            upload_date=datetime.fromisoformat(row[6]) if row[6] else datetime.now(),
            text_content=row[7] or "",
            chunks=chunks,
            metadata=metadata
        )
        
        return processed_file
        
    except Exception as e:
        logger.error(f"[FILES] Error getting file {file_id}: {e}")
        return None
    finally:
        conn.close()

def delete_file(file_id: str, username: str) -> bool:
    """Löscht Datei mit Berechtigung-Check"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Berechtigung prüfen
        cursor.execute("SELECT id FROM uploaded_files WHERE id = ? AND username = ?", 
                      (file_id, username))
        if not cursor.fetchone():
            return False
        
        # Chunks löschen
        cursor.execute("DELETE FROM file_chunks WHERE file_id = ?", (file_id,))
        
        # Usage-Logs löschen
        cursor.execute("DELETE FROM file_usage WHERE file_id = ?", (file_id,))
        
        # Haupteintrag löschen
        cursor.execute("DELETE FROM uploaded_files WHERE id = ?", (file_id,))
        
        # Physische Datei löschen
        file_path = get_file_path(file_id)
        if file_path.exists():
            file_path.unlink()
        
        conn.commit()
        logger.info(f"[FILES] File deleted: {file_id}")
        return True
        
    except Exception as e:
        logger.error(f"[FILES] Error deleting file {file_id}: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def log_file_usage(file_id: str, username: str, action: str, cursor=None):
    """Loggt File-Usage"""
    close_conn = False
    if cursor is None:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        close_conn = True
    
    try:
        cursor.execute("""
            INSERT INTO file_usage (file_id, username, action)
            VALUES (?, ?, ?)
        """, (file_id, username, action))
        
        if close_conn:
            conn.commit()
            
    except Exception as e:
        logger.error(f"[FILES] Usage logging error: {e}")
    finally:
        if close_conn:
            conn.close()

# ──────────────────────────────
# File Search and Context Integration
# ──────────────────────────────

def search_in_files(username: str, query: str, file_id: Optional[str] = None, 
                   limit: int = 10) -> List[Dict]:
    """Durchsucht Dateiinhalte"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        if file_id:
            # Suche in spezifischer Datei
            cursor.execute("""
                SELECT fc.file_id, fc.content, fc.chunk_index,
                       uf.original_name, uf.file_type, uf.thread_id
                FROM file_chunks fc
                JOIN uploaded_files uf ON fc.file_id = uf.id
                WHERE fc.content LIKE ? AND uf.id = ? AND uf.username = ?
                ORDER BY fc.chunk_index
                LIMIT ?
            """, (f"%{query}%", file_id, username, limit))
        else:
            # Suche in allen Dateien
            cursor.execute("""
                SELECT fc.file_id, fc.content, fc.chunk_index,
                       uf.original_name, uf.file_type, uf.thread_id
                FROM file_chunks fc
                JOIN uploaded_files uf ON fc.file_id = uf.id
                WHERE fc.content LIKE ? AND uf.username = ?
                ORDER BY uf.upload_date DESC, fc.chunk_index
                LIMIT ?
            """, (f"%{query}%", username, limit))
        
        rows = cursor.fetchall()
        
        results = []
        for row in rows:
            # Highlight Search-Term
            highlighted_content = highlight_search_term(row[1], query, max_length=300)
            
            results.append({
                'file_id': row[0],
                'content': row[1],
                'highlighted_content': highlighted_content,
                'chunk_index': row[2],
                'file_name': row[3],
                'file_type': row[4],
                'thread_id': row[5]
            })
        
        return results
        
    except sqlite3.OperationalError:
        return []
    except Exception as e:
        logger.error(f"[FILES] File search error: {e}")
        return []
    finally:
        conn.close()

def get_file_context_for_chat(username: str, thread_id: str) -> List[Dict]:
    """Holt Dateien-Kontext für Chat"""
    files = get_user_files(username, thread_id)
    
    context_files = []
    for file in files:
        if file['chunks_count'] > 0:  # Nur Dateien mit Text-Content
            context_files.append({
                'id': file['id'],
                'name': file['original_name'],
                'type': file['file_type'],
                'summary': file['summary'][:200] + "..." if len(file['summary']) > 200 else file['summary'],
                'chunks_count': file['chunks_count']
            })
    
    return context_files

def get_relevant_file_chunks(username: str, query: str, thread_id: Optional[str] = None, 
                           max_chunks: int = 3) -> str:
    """Holt relevante File-Chunks für AI-Context"""
    search_results = search_in_files(username, query, limit=max_chunks * 2)
    
    if not search_results:
        return ""
    
    # Thread-Filter falls gesetzt
    if thread_id:
        search_results = [r for r in search_results if r['thread_id'] == thread_id]
    
    # Beste Ergebnisse auswählen
    relevant_chunks = []
    seen_files = set()
    
    for result in search_results[:max_chunks]:
        if result['file_id'] not in seen_files:
            relevant_chunks.append(f"[Aus {result['file_name']}]: {result['content'][:500]}...")
            seen_files.add(result['file_id'])
    
    if relevant_chunks:
        return "\n\n--- Relevante Dokumente ---\n" + "\n\n".join(relevant_chunks) + "\n--- Ende Dokumente ---\n"
    
    return ""

# ──────────────────────────────
# File Statistics and Analytics
# ──────────────────────────────

def get_user_file_stats(username: str) -> Dict:
    """Holt File-Statistiken für User"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Grundstatistiken
        cursor.execute("""
            SELECT 
                COUNT(*) as total_files,
                SUM(file_size) as total_size,
                SUM(chunks_count) as total_chunks,
                AVG(file_size) as avg_file_size
            FROM uploaded_files 
            WHERE username = ?
        """, (username,))
        
        stats = cursor.fetchone()
        
        # Dateityp-Verteilung
        cursor.execute("""
            SELECT file_type, COUNT(*) as count, SUM(file_size) as size
            FROM uploaded_files 
            WHERE username = ?
            GROUP BY file_type
            ORDER BY count DESC
        """, (username,))
        
        type_distribution = cursor.fetchall()
        
        # Neueste Uploads
        cursor.execute("""
            SELECT original_name, file_type, upload_date, file_size
            FROM uploaded_files 
            WHERE username = ?
            ORDER BY upload_date DESC 
            LIMIT 5
        """, (username,))
        
        recent_files = cursor.fetchall()
        
        # Usage-Statistiken
        cursor.execute("""
            SELECT action, COUNT(*) as count
            FROM file_usage fu
            JOIN uploaded_files uf ON fu.file_id = uf.id
            WHERE uf.username = ?
            GROUP BY action
        """, (username,))
        
        usage_stats = cursor.fetchall()
        
        return {
            'total_files': stats[0] or 0,
            'total_size_mb': round((stats[1] or 0) / (1024 * 1024), 2),
            'total_chunks': stats[2] or 0,
            'avg_file_size_mb': round((stats[3] or 0) / (1024 * 1024), 2),
            'type_distribution': [
                {
                    'type': row[0],
                    'count': row[1],
                    'size_mb': round(row[2] / (1024 * 1024), 2)
                }
                for row in type_distribution
            ],
            'recent_files': [
                {
                    'name': row[0],
                    'type': row[1],
                    'date': row[2],
                    'size_mb': round(row[3] / (1024 * 1024), 2)
                }
                for row in recent_files
            ],
            'usage_stats': dict(usage_stats)
        }
        
    except sqlite3.OperationalError:
        return {
            'total_files': 0,
            'total_size_mb': 0,
            'total_chunks': 0,
            'avg_file_size_mb': 0,
            'type_distribution': [],
            'recent_files': [],
            'usage_stats': {}
        }
    finally:
        conn.close()

def get_global_file_stats() -> Dict:
    """Globale File-Statistiken für Admin"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Gesamtstatistiken
        cursor.execute("""
            SELECT 
                COUNT(*) as total_files,
                COUNT(DISTINCT username) as unique_users,
                SUM(file_size) as total_size,
                AVG(file_size) as avg_file_size
            FROM uploaded_files
        """)
        
        global_stats = cursor.fetchone()
        
        # Top File-Typen
        cursor.execute("""
            SELECT file_type, COUNT(*) as count, SUM(file_size) as total_size
            FROM uploaded_files
            GROUP BY file_type
            ORDER BY count DESC
        """)
        
        file_types = cursor.fetchall()
        
        # Upload-Aktivität (letzte 30 Tage)
        cursor.execute("""
            SELECT DATE(upload_date) as date, COUNT(*) as uploads
            FROM uploaded_files 
            WHERE upload_date >= datetime('now', '-30 days')
            GROUP BY DATE(upload_date)
            ORDER BY date DESC
        """)
        
        daily_uploads = cursor.fetchall()
        
        # Größte Files
        cursor.execute("""
            SELECT original_name, username, file_size, upload_date
            FROM uploaded_files
            ORDER BY file_size DESC
            LIMIT 10
        """)
        
        largest_files = cursor.fetchall()
        
        return {
            'total_files': global_stats[0] or 0,
            'unique_users': global_stats[1] or 0,
            'total_size_gb': round((global_stats[2] or 0) / (1024 * 1024 * 1024), 2),
            'avg_file_size_mb': round((global_stats[3] or 0) / (1024 * 1024), 2),
            'file_types': [
                {
                    'type': row[0],
                    'count': row[1],
                    'size_gb': round(row[2] / (1024 * 1024 * 1024), 2)
                }
                for row in file_types
            ],
            'daily_uploads': [
                {'date': row[0], 'count': row[1]}
                for row in daily_uploads
            ],
            'largest_files': [
                {
                    'name': row[0],
                    'username': row[1],
                    'size_mb': round(row[2] / (1024 * 1024), 2),
                    'date': row[3]
                }
                for row in largest_files
            ]
        }
        
    except sqlite3.OperationalError:
        return {
            'total_files': 0,
            'unique_users': 0,
            'total_size_gb': 0,
            'avg_file_size_mb': 0,
            'file_types': [],
            'daily_uploads': [],
            'largest_files': []
        }
    finally:
        conn.close()

# ──────────────────────────────
# Enhanced Chat Integration
# ──────────────────────────────

async def get_ai_response_with_files(current_message: str, chat_history: list, 
                                   username: str, thread_id: str, persona: str = "standard",
                                   model_override: str = None, include_files: bool = True) -> Dict:
    """
    Erweiterte AI-Response mit File-Context
    """
    # File-Context hinzufügen falls gewünscht
    enhanced_message = current_message
    
    if include_files:
        # Relevante File-Chunks basierend auf User-Message holen
        file_context = get_relevant_file_chunks(username, current_message, thread_id, max_chunks=2)
        
        if file_context:
            enhanced_message = f"{current_message}\n{file_context}"
    
    # Normale AI-Response mit enhanced message
    return await get_ai_response_with_model(
        current_message=enhanced_message,
        chat_history=chat_history,
        username=username,
        thread_id=thread_id,
        persona=persona,
        model_override=model_override
    )

# ──────────────────────────────
# Background Tasks für File System
# ──────────────────────────────

async def cleanup_orphaned_files():
    """Bereinigt verwaiste Dateien"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # Dateien ohne DB-Eintrag finden
        upload_path = ensure_upload_directory()
        physical_files = set(f.stem for f in upload_path.glob("*.dat"))
        
        cursor.execute("SELECT id FROM uploaded_files")
        db_files = set(row[0] for row in cursor.fetchall())
        
        orphaned_files = physical_files - db_files
        
        # Verwaiste Dateien löschen
        deleted_count = 0
        for file_id in orphaned_files:
            file_path = get_file_path(file_id)
            if file_path.exists():
                file_path.unlink()
                deleted_count += 1
        
        # Verwaiste DB-Einträge finden
        missing_files = db_files - physical_files
        
        # Verwaiste DB-Einträge löschen
        for file_id in missing_files:
            cursor.execute("DELETE FROM file_chunks WHERE file_id = ?", (file_id,))
            cursor.execute("DELETE FROM file_usage WHERE file_id = ?", (file_id,))
            cursor.execute("DELETE FROM uploaded_files WHERE id = ?", (file_id,))
        
        conn.commit()
        conn.close()
        
        return {
            'deleted_orphaned_files': deleted_count,
            'deleted_orphaned_db_entries': len(missing_files),
            'cleanup_date': datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"File cleanup error: {e}")
        raise

async def update_file_statistics():
    """Aktualisiert File-Statistiken"""
    try:
        global_stats = get_global_file_stats()
        
        # Speicher-Warnungen
        warnings = []
        if global_stats['total_size_gb'] > 10:
            warnings.append(f"Hoher Speicherverbrauch: {global_stats['total_size_gb']} GB")
        
        if global_stats['total_files'] > 10000:
            warnings.append(f"Viele Dateien: {global_stats['total_files']}")
        
        return {
            'total_files': global_stats['total_files'],
            'total_size_gb': global_stats['total_size_gb'],
            'unique_users': global_stats['unique_users'],
            'warnings': warnings,
            'update_time': datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"File statistics update error: {e}")
        raise

def setup_file_background_tasks():
    """Registriert File-Processing Background-Tasks"""
    task_manager.register_task_function(
        "cleanup_orphaned_files",
        cleanup_orphaned_files,
        "Bereinigt verwaiste Dateien ohne DB-Referenz"
    )
    
    task_manager.register_task_function(
        "update_file_statistics",
        update_file_statistics,
        "Aktualisiert File-System-Statistiken"
    )
    
    # Cleanup täglich
    task_manager.schedule_recurring_task(
        "cleanup_files_scheduled",
        cleanup_orphaned_files,
        interval_seconds=86400,  # Täglich
        run_immediately=False
    )
    
    # Statistiken alle 6 Stunden
    task_manager.schedule_recurring_task(
        "file_stats_scheduled",
        update_file_statistics,
        interval_seconds=21600,  # 6 Stunden
        run_immediately=True
    )

# ──────────────────────────────
# Cache Functions für Files
# ──────────────────────────────

@cache_result("user_files", ttl=300)
def get_user_files_cached(username: str, thread_id: str = None):
    """Cached Version von get_user_files"""
    return get_user_files(username, thread_id)

@cache_result("file_stats", ttl=600)
def get_user_file_stats_cached(username: str):
    """Cached Version von get_user_file_stats"""
    return get_user_file_stats(username)

def invalidate_file_cache(username: str):
    """File-bezogene Cache-Einträge löschen"""
    keys_to_delete = [
        f"user_files:{username}",
        f"user_files:{username}:None",
        f"file_stats:{username}"
    ]
    
    for key in keys_to_delete:
        app_cache.delete(key)

# Rate-Limiting Funktionen
def load_rate_limits():
    """Lädt Rate-Limit-Daten"""
    if os.path.exists(RATE_LIMIT_FILE):
        try:
            with open(RATE_LIMIT_FILE, 'r') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_rate_limits(limits):
    """Speichert Rate-Limit-Daten"""
    with open(RATE_LIMIT_FILE, 'w') as f:
        json.dump(limits, f)

def check_rate_limit(username: str) -> bool:
    """Prüft ob User Rate-Limit erreicht hat"""
    limits = load_rate_limits()
    current_time = _pytime.time()
    
    user_data = limits.get(username, {"messages": [], "last_reset": current_time})
    
    # Alte Nachrichten entfernen (älter als 1 Stunde)
    hour_ago = current_time - 3600
    user_data["messages"] = [msg_time for msg_time in user_data["messages"] if msg_time > hour_ago]
    
    # Prüfen ob Limit erreicht
    if len(user_data["messages"]) >= MESSAGES_PER_HOUR:
        return False
    
    # Neue Nachricht hinzufügen
    user_data["messages"].append(current_time)
    limits[username] = user_data
    save_rate_limits(limits)
    
    return True

# Session-Management
def update_session_activity(request: Request):
    """Aktualisiert die letzte Aktivität in der Session"""
    request.session["last_activity"] = _pytime.time()

def check_session_timeout(request: Request) -> bool:
    """Prüft ob Session abgelaufen ist"""
    last_activity = request.session.get("last_activity")
    if not last_activity:
        return True
    
    return (_pytime.time() - last_activity) > SESSION_TIMEOUT_SECONDS

def require_active_session(request: Request):
    """Middleware-ähnliche Funktion für Session-Check"""
    if not request.session.get("username"):
        return RedirectResponse("/", status_code=302)
    
    if check_session_timeout(request):
        request.session.clear()
        return RedirectResponse("/?timeout=1", status_code=302)
    
    update_session_activity(request)
    return None

def is_admin(request: Request) -> bool:
    username = request.session.get("username")
    if username:
        user = get_user(username)
        return user and user["is_admin"]
    return False

def admin_redirect_guard(request: Request):
    if not is_admin(request):
        return RedirectResponse("/", status_code=302)
    return None

# ──────────────────────────────
# Routes - Auth
# ──────────────────────────────
@app.get("/", response_class=HTMLResponse)
async def service_selection_page(request: Request):
    """Service-Auswahl Homepage"""
    request.session.clear()
    return templates.TemplateResponse("service_selection.html", {
        "request": request
    })

@app.post("/login", response_class=HTMLResponse)
async def login(request: Request, username: str = Form(...), password: str = Form(...)):
    if check_login(username, password):
        user = get_user(username)
        if user["is_blocked"]:
            return templates.TemplateResponse("login.html", {
                "request": request,
                "error": "Du wurdest vom Admin gesperrt."
            })
        
        request.session["username"] = username
        update_session_activity(request)
        
        if user["is_admin"]:
            return RedirectResponse("/admin", status_code=302)
        else:
            return RedirectResponse("/chat", status_code=302)
    
    return templates.TemplateResponse("login.html", {
        "request": request, 
        "error": "Falsche Anmeldedaten"
    })

@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    return templates.TemplateResponse("register.html", {"request": request})

@app.post("/register", response_class=HTMLResponse)
async def register(request: Request, 
                  username: str = Form(...), 
                  password: str = Form(...), 
                  question: str = Form(...), 
                  answer: str = Form(...)):
    
    if save_user(username, password, question, answer):
        return RedirectResponse("/", status_code=302)
    
    return templates.TemplateResponse("register.html", {
        "request": request, 
        "error": "Benutzer existiert bereits"
    })

@app.get("/reset", response_class=HTMLResponse)
async def reset_page(request: Request):
    return templates.TemplateResponse("reset.html", {
        "request": request, 
        "error": "", 
        "success": ""
    })

@app.post("/reset", response_class=HTMLResponse)
async def reset_post(request: Request, 
                    username: str = Form(...), 
                    answer: str = Form(...), 
                    new_password: str = Form(...)):
    
    if verify_security_answer(username, answer):
        reset_password(username, new_password)
        return RedirectResponse("/", status_code=302)
    
    return templates.TemplateResponse("reset.html", {
        "request": request, 
        "error": "Antwort falsch", 
        "success": ""
    })

@app.get("/logout")
async def logout(request: Request):
    request.session.clear()
    return RedirectResponse("/", status_code=302)

@app.get("/login", response_class=HTMLResponse)
async def login_with_service(request: Request, service: str = None):
    """Login mit Service-Parameter"""
    if service:
        request.session["selected_service"] = service
    
    return templates.TemplateResponse("login.html", {
        "request": request,
        "service": service
    })
# ==================== GASTRO-SPEZIFISCHER LOGIN ====================

@app.get("/gastro-login", response_class=HTMLResponse)
async def gastro_login_page(request: Request):
    """Separate Login-Seite für Gastro-Service"""
    return templates.TemplateResponse("gastro_login.html", {
        "request": request
    })

@app.post("/gastro-login")
async def gastro_login_post(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    db: Session = Depends(get_db)
):
    """Login mit Weiterleitung zu Gastro-Dashboard"""
    user = authenticate_user(db, username, password)
    
    if not user:
        return templates.TemplateResponse("gastro_login.html", {
            "request": request,
            "error": "Ungültige Anmeldedaten"
        })
    
    request.session["user_id"] = user.id
    return RedirectResponse("/gastro-dashboard", status_code=303)

@app.get("/gastro-dashboard", response_class=HTMLResponse)
async def gastro_dashboard(request: Request):
    """Gastro-spezifisches Dashboard"""
    redirect = require_active_session(request)
    if redirect:
        return RedirectResponse("/gastro-login", status_code=303)
    
    username = request.session.get("username")
    
    return templates.TemplateResponse("gastro_dashboard.html", {
        "request": request,
        "username": username
    })
# ──────────────────────────────
# Routes - Chat
# ──────────────────────────────
# ──────────────────────────────
# Updated Chat Routes mit Threading-Support
# ──────────────────────────────

@app.get("/chat", response_class=HTMLResponse)
async def chat_page_with_threading(request: Request):
    """Chat-Seite mit Threading-Support"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    # Aktuellen Thread aus URL oder Session holen
    thread_id = request.query_params.get('thread', get_current_thread_id(request))
    set_current_thread_id(request, thread_id)
    
    # Threads des Users holen
    threads = get_user_threads(username)
    
    # Aktuellen Thread validieren
    thread_exists = any(t['id'] == thread_id for t in threads)
    if not thread_exists and thread_id != 'default':
        thread_id = 'default'
        set_current_thread_id(request, thread_id)
    
    # Chat-Verlauf für aktuellen Thread
    history = get_thread_history(username, thread_id)
    
    # Thread-Info
    current_thread = get_thread_info(username, thread_id) or {
        'id': 'default',
        'title': 'Haupt-Unterhaltung',
        'message_count': len(history)
    }
    
    return templates.TemplateResponse("chat_with_threads.html", {
        "request": request,
        "username": username,
        "chat_history": history,
        "threads": threads,
        "current_thread": current_thread,
        "session_timeout_minutes": SESSION_TIMEOUT_MINUTES
    })

# ──────────────────────────────
# File Processing Routes
# ──────────────────────────────

from fastapi import UploadFile, File
from fastapi.responses import FileResponse

@app.get("/files", response_class=HTMLResponse)
async def files_page(request: Request):
    """File-Management-Seite"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    thread_id = request.query_params.get('thread')
    
    # User-Dateien holen
    user_files = get_user_files_cached(username, thread_id)
    
    # File-Statistiken
    file_stats = get_user_file_stats_cached(username)
    
    # Thread-Info falls spezifisch
    current_thread = None
    if thread_id:
        current_thread = get_thread_info(username, thread_id)
    
    # Threads für Filter
    threads = get_user_threads(username)
    
    return templates.TemplateResponse("files.html", {
        "request": request,
        "username": username,
        "user_files": user_files,
        "file_stats": file_stats,
        "current_thread": current_thread,
        "threads": threads,
        "max_file_size_mb": MAX_FILE_SIZE // (1024 * 1024),
        "max_files_per_user": MAX_FILES_PER_USER
    })

@app.post("/files/upload")
async def upload_file(request: Request, 
                     file: UploadFile = File(...),
                     thread_id: str = Form(None)):
    """File-Upload-Endpoint"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    try:
        # File-Content lesen
        file_content = await file.read()
        filename = file.filename or "unbekannt"
        
        # Validierung
        is_valid, error_message = validate_file_upload(file_content, filename, username)
        if not is_valid:
            return RedirectResponse(f"/files?error={error_message}", status_code=302)
        
        # File verarbeiten
        processed_file = process_uploaded_file(
            file_content=file_content,
            filename=filename,
            username=username,
            thread_id=thread_id
        )
        
        # In DB speichern
        if save_processed_file_to_db(processed_file):
            # Cache invalidieren
            invalidate_file_cache(username)
            
            # Usage loggen
            log_file_usage(processed_file.id, username, "upload")
            
            logger.info(f"[FILES] File uploaded: {filename} by {username}")
            return RedirectResponse(f"/files?success=uploaded&file_id={processed_file.id}", status_code=302)
        else:
            return RedirectResponse("/files?error=save_failed", status_code=302)
        
    except Exception as e:
        logger.error(f"[FILES] Upload error: {e}")
        return RedirectResponse(f"/files?error=processing_failed", status_code=302)

@app.get("/files/{file_id}")
async def view_file(request: Request, file_id: str):
    """File-Detail-Ansicht"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    # File laden mit Berechtigung-Check
    processed_file = get_file_by_id(file_id, username)
    if not processed_file:
        return RedirectResponse("/files?error=file_not_found", status_code=302)
    
    # Usage loggen
    log_file_usage(file_id, username, "view")
    
    return templates.TemplateResponse("file_detail.html", {
        "request": request,
        "username": username,
        "file": processed_file
    })

@app.post("/files/{file_id}/delete")
async def delete_file_route(request: Request, file_id: str):
    """File löschen"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    if delete_file(file_id, username):
        invalidate_file_cache(username)
        logger.info(f"[FILES] File deleted: {file_id} by {username}")
        return RedirectResponse("/files?success=deleted", status_code=302)
    else:
        return RedirectResponse("/files?error=delete_failed", status_code=302)

@app.get("/files/{file_id}/download")
async def download_file(request: Request, file_id: str):
    """File-Download"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    # File-Info laden
    processed_file = get_file_by_id(file_id, username)
    if not processed_file:
        raise HTTPException(status_code=404, detail="File not found")
    
    # Physische Datei prüfen
    file_path = get_file_path(file_id)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Physical file not found")
    
    # Usage loggen
    log_file_usage(file_id, username, "download")
    
    return FileResponse(
        path=file_path,
        filename=processed_file.original_name,
        media_type='application/octet-stream'
    )

@app.get("/api/files/search")
async def api_search_files(request: Request):
    """API für File-Suche"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    query = request.query_params.get('q', '').strip()
    file_id = request.query_params.get('file_id')
    limit = min(int(request.query_params.get('limit', 10)), 50)
    
    if not query:
        return {"error": "Suchbegriff fehlt"}
    
    # File-Suche
    results = search_in_files(username, query, file_id, limit)
    
    return {
        "query": query,
        "total_results": len(results),
        "results": results
    }

@app.get("/api/files/context/{thread_id}")
async def api_files_context(request: Request, thread_id: str):
    """API für File-Context in Thread"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    
    # File-Context für Thread
    context_files = get_file_context_for_chat(username, thread_id)
    
    return {
        "thread_id": thread_id,
        "files": context_files,
        "total_files": len(context_files)
    }

# ──────────────────────────────
# Enhanced Chat Route mit File-Context
# ──────────────────────────────

@app.post("/chat")
async def chat_with_files_and_models(req: Request):
    """Chat mit Multi-AI und File-Context Support"""
    redirect = require_active_session(req)
    if redirect:
        return {"reply": "Session abgelaufen. Bitte neu anmelden.", "redirect": "/"}
    
    username = req.session.get("username")
    
    # Rate-Limit-Prüfung
    rate_limit_result = check_enhanced_rate_limit(username)
    if not rate_limit_result["allowed"]:
        tier = rate_limit_result.get("tier", "free")
        tier_name = SUBSCRIPTION_TIERS.get(tier, {}).get("name", tier)
        
        return {
            "reply": f"Rate-Limit erreicht für {tier_name}-Plan!",
            "rate_limit": True,
            "tier": tier
        }
    
    data = await req.json()
    user_message = data.get("message", "")
    thread_id = data.get("thread_id", get_current_thread_id(req))
    model_override = data.get("model_id")
    include_files = data.get("include_files", True)  # File-Context aktiviert?
    
    if not user_message.strip():
        return {"reply": "Leere Nachricht."}
    
    # Thread validieren/erstellen
    if not get_thread_info(username, thread_id):
        if thread_id == 'default':
            conn = sqlite3.connect(DB_PATH)
            cursor = conn.cursor()
            cursor.execute("""
                INSERT OR IGNORE INTO chat_threads (id, username, title)
                VALUES ('default', ?, 'Haupt-Unterhaltung')
            """, (username,))
            conn.commit()
            conn.close()
        else:
            thread_id = 'default'
    
    # Persona prüfen
    current_persona = get_user_persona_cached(username)
    available_personas = get_available_personas_for_user_cached(username)
    
    if current_persona not in available_personas:
        current_persona = "standard"
        save_user_persona_cached(username, current_persona)
    
    # User-Nachricht speichern
    save_message_to_thread(username, thread_id, "user", user_message)
    
    # Chat-Historie laden
    history = get_thread_history(username, thread_id)
    
    try:
        # AI-Response mit File-Context (falls aktiviert)
        ai_result = await get_ai_response_with_files(
            current_message=user_message,
            chat_history=history,
            username=username,
            thread_id=thread_id,
            persona=current_persona,
            model_override=model_override,
            include_files=include_files
        )
        
        raw_response = ai_result['content']
        model_info = ai_result['model_info']
        
        # Markdown rendern
        rendered_response = render_markdown_simple(raw_response)
        
        # AI-Antwort mit Model-Info speichern
        save_message_to_thread_with_model(
            username, thread_id, "assistant", raw_response, model_info
        )
        
        # Auto-Titel für neue Threads
        thread_info = get_thread_info(username, thread_id)
        if thread_info and thread_info['message_count'] <= 2 and thread_id != 'default':
            update_thread_title_from_first_message(username, thread_id)
        
        # Caches invalidieren
        invalidate_thread_cache(username)
        invalidate_model_cache(username)
        invalidate_search_on_new_message(username, thread_id)
        
        # File-Context-Info für Response
        file_context_info = None
        if include_files:
            context_files = get_file_context_for_chat(username, thread_id)
            if context_files:
                file_context_info = {
                    "files_available": len(context_files),
                    "files_used": True  # Könnte verfeinert werden
                }
        
        return {
            "reply": rendered_response,
            "raw_reply": raw_response,
            "thread_id": thread_id,
            "model_info": {
                "model_used": model_info['model_used'],
                "provider": model_info['provider'],
                "tokens_used": model_info['tokens_used'],
                "response_time": model_info['duration'],
                "cost": model_info['cost'],
                "is_fallback": model_info.get('is_fallback', False)
            },
            "file_context": file_context_info,
            "remaining_messages": {
                "hour": rate_limit_result["remaining_hour"],
                "day": rate_limit_result["remaining_day"]
            }
        }
        
    except Exception as e:
        logger.error(f"Enhanced chat error for {username}: {str(e)}")
        return {"reply": "Ein Fehler ist aufgetreten. Versuche es erneut."}

# ──────────────────────────────
# Admin File Management
# ──────────────────────────────

@app.get("/admin/files", response_class=HTMLResponse)
async def admin_files_page(request: Request):
    """Admin-Seite für File-Management"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    # Globale File-Statistiken
    global_stats = get_global_file_stats()
    
    return templates.TemplateResponse("admin_files.html", {
        "request": request,
        "global_stats": global_stats,
        "max_file_size_mb": MAX_FILE_SIZE // (1024 * 1024),
        "max_files_per_user": MAX_FILES_PER_USER
    })

@app.get("/admin/files/{username}")
async def admin_user_files(request: Request, username: str):
    """Admin-Ansicht für User-Files"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    user_files = get_user_files(username)
    user_stats = get_user_file_stats(username)
    
    return templates.TemplateResponse("admin_user_files.html", {
        "request": request,
        "target_username": username,
        "user_files": user_files,
        "user_stats": user_stats
    })

@app.post("/admin/files/{file_id}/delete")
async def admin_delete_file(request: Request, file_id: str):
    """Admin löscht User-File"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    # File-Info holen (ohne User-Berechtigung)
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        cursor.execute("SELECT username FROM uploaded_files WHERE id = ?", (file_id,))
        row = cursor.fetchone()
        
        if not row:
            return {"success": False, "message": "File nicht gefunden"}
        
        file_username = row[0]
        
        # File löschen
        if delete_file(file_id, file_username):
            logger.info(f"[ADMIN] File deleted by admin: {file_id}")
            return {"success": True, "message": "File gelöscht"}
        else:
            return {"success": False, "message": "Löschung fehlgeschlagen"}
            
    finally:
        conn.close()

@app.post("/admin/files/cleanup")
async def admin_cleanup_files(request: Request):
    """Admin startet File-Cleanup"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    # Cleanup als Background-Task starten
    task_id = await task_manager.run_task(
        "manual_file_cleanup",
        cleanup_orphaned_files
    )
    
    logger.info(f"[ADMIN] Manual file cleanup started (Task ID: {task_id})")
    return RedirectResponse("/admin/files?cleanup_started=1", status_code=302)

# ──────────────────────────────
# File-Integration in Chat-Template
# ──────────────────────────────

@app.get("/api/chat/files/{thread_id}")
async def api_chat_files_info(request: Request, thread_id: str):
    """File-Info für Chat-Interface"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    
    # Dateien im Thread
    thread_files = get_user_files(username, thread_id)
    
    # Kompakte Info für Chat-UI
    files_info = []
    for file in thread_files[:10]:  # Max 10 für UI
        files_info.append({
            "id": file['id'],
            "name": file['original_name'],
            "type": file['file_type'],
            "size_mb": round(file['file_size'] / (1024 * 1024), 2),
            "summary": file['summary'][:100] + "..." if len(file['summary']) > 100 else file['summary']
        })
    
    return {
        "thread_id": thread_id,
        "files": files_info,
        "total_files": len(thread_files),
        "has_files": len(thread_files) > 0
    }

# ──────────────────────────────
# Admin Model Management
# ──────────────────────────────

@app.get("/admin/models", response_class=HTMLResponse)
async def admin_models_page(request: Request):
    """Admin-Seite für Model-Management"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    # Globale Model-Statistiken
    global_stats = get_global_model_stats()
    
    # Provider-Status
    provider_status = {}
    for provider in ModelProvider:
        provider_obj = model_manager.providers.get(provider)
        provider_status[provider.value] = {
            'available': provider_obj.is_available if provider_obj else False,
            'models_count': len([m for m in AI_MODELS.values() if m.provider == provider and m.is_active])
        }
    
    # Model-Konfiguration
    models_config = {}
    for model_id, model in AI_MODELS.items():
        models_config[model_id] = {
            'display_name': model.display_name,
            'provider': model.provider.value,
            'is_active': model.is_active,
            'subscription_tier': model.subscription_tier,
            'cost_per_1k_tokens': model.cost_per_1k_tokens
        }
    
    return templates.TemplateResponse("admin_models.html", {
        "request": request,
        "global_stats": global_stats,
        "provider_status": provider_status,
        "models_config": models_config
    })

@app.post("/admin/models/test-providers")
async def admin_test_providers(request: Request):
    """Testet alle Provider manuell"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    # Test als Background-Task starten
    task_id = await task_manager.run_task(
        "manual_provider_test",
        test_model_providers
    )
    
    logger.info(f"[ADMIN] Manual provider test started (Task ID: {task_id})")
    return RedirectResponse("/admin/models?test_started=1", status_code=302)

@app.post("/admin/models/{model_id}/toggle")
async def admin_toggle_model(request: Request, model_id: str):
    """Aktiviert/Deaktiviert ein Model (Admin)"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    if model_id in AI_MODELS:
        AI_MODELS[model_id].is_active = not AI_MODELS[model_id].is_active
        status = "aktiviert" if AI_MODELS[model_id].is_active else "deaktiviert"
        
        logger.info(f"[ADMIN] Model {model_id} {status}")
        return {"success": True, "message": f"Model {status}"}
    else:
        return {"success": False, "message": "Model nicht gefunden"}

@app.get("/admin/models/usage/{username}")
async def admin_user_model_usage(request: Request, username: str):
    """Admin-Ansicht für User-Model-Usage"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    user_stats = get_user_model_stats(username)
    
    return templates.TemplateResponse("admin_user_models.html", {
        "request": request,
        "target_username": username,
        "user_stats": user_stats
    })

# ──────────────────────────────
# Model Integration in Chat-Template
# ──────────────────────────────

@app.get("/api/chat/models")
async def api_chat_models_info(request: Request):
    """Model-Info für Chat-Interface"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    thread_id = request.query_params.get('thread', get_current_thread_id(request))
    
    # Aktuelles Model für Thread
    thread_model = get_thread_preferred_model(username, thread_id)
    user_model = get_user_preferred_model_cached(username)
    effective_model = thread_model or user_model
    
    # Verfügbare Modelle
    available_models = get_available_models_for_user_cached(username)
    
    return {
        "current_model": effective_model,
        "thread_specific": thread_model is not None,
        "available_models": {
            model_id: {
                "name": model.display_name,
                "personality": model.personality,
                "provider": model.provider.value
            }
            for model_id, model in available_models.items()
        }
    }

# ──────────────────────────────
# Cost Tracking und Billing
# ──────────────────────────────

@app.get("/api/models/costs")
async def api_model_costs(request: Request):
    """API für Model-Kosten des Users"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # Kosten der letzten 30 Tage
        cursor.execute("""
            SELECT DATE(timestamp) as date, 
                   SUM(cost) as daily_cost,
                   COUNT(*) as requests
            FROM model_usage 
            WHERE username = ? AND timestamp >= datetime('now', '-30 days')
            GROUP BY DATE(timestamp)
            ORDER BY date DESC
        """, (username,))
        
        daily_costs = cursor.fetchall()
        
        # Gesamtkosten pro Model
        cursor.execute("""
            SELECT model_name, SUM(cost) as total_cost, COUNT(*) as usage_count
            FROM model_usage 
            WHERE username = ?
            GROUP BY model_name
            ORDER BY total_cost DESC
        """, (username,))
        
        model_costs = cursor.fetchall()
        
        conn.close()
        
        return {
            "daily_costs": [
                {"date": row[0], "cost": round(row[1], 4), "requests": row[2]}
                for row in daily_costs
            ],
            "model_costs": [
                {"model": row[0], "total_cost": round(row[1], 4), "usage_count": row[2]}
                for row in model_costs
            ],
            "total_cost": sum(row[1] for row in model_costs)
        }
        
    except sqlite3.OperationalError:
        return {"daily_costs": [], "model_costs": [], "total_cost": 0}

# ──────────────────────────────
# Chat Search Routes
# ──────────────────────────────

@app.get("/search", response_class=HTMLResponse)
async def search_page(request: Request):
    """Chat-Suche Hauptseite"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    # User-Threads für Thread-Filter holen
    threads = get_user_threads(username)
    
    # Such-Statistiken
    stats = get_search_statistics_cached(username)
    
    return templates.TemplateResponse("search.html", {
        "request": request,
        "username": username,
        "threads": threads,
        "stats": stats
    })

@app.get("/api/search")
async def api_search_messages(request: Request):
    """API-Endpoint für Chat-Suche"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    
    # Query-Parameter
    query = request.query_params.get('q', '').strip()
    thread_id = request.query_params.get('thread', None)
    days_back = request.query_params.get('days', None)
    limit = min(int(request.query_params.get('limit', 50)), 100)  # Max 100
    
    if not query:
        return {"error": "Suchbegriff fehlt"}
    
    if len(query) < 2:
        return {"error": "Suchbegriff zu kurz (min. 2 Zeichen)"}
    
    try:
        days_back = int(days_back) if days_back else None
    except ValueError:
        days_back = None
    
    # Suche ausführen
    results = search_chat_messages(
        username=username,
        query=query,
        thread_id=thread_id,
        days_back=days_back,
        limit=limit
    )
    
    return {
        "query": query,
        "total_results": len(results),
        "results": results,
        "filters": {
            "thread_id": thread_id,
            "days_back": days_back,
            "limit": limit
        }
    }

@app.get("/api/search/advanced")
async def api_advanced_search(request: Request):
    """API-Endpoint für erweiterte Suche"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    
    # Query-Parameter
    query = request.query_params.get('q', '').strip()
    role = request.query_params.get('role', None)
    thread_id = request.query_params.get('thread', None)
    date_from = request.query_params.get('from', None)
    date_to = request.query_params.get('to', None)
    limit = min(int(request.query_params.get('limit', 50)), 100)
    
    if not query:
        return {"error": "Suchbegriff fehlt"}
    
    # Filter zusammenstellen
    filters = {
        'role': role,
        'thread_id': thread_id,
        'date_from': date_from,
        'date_to': date_to,
        'limit': limit
    }
    
    # Erweiterte Suche ausführen
    results = advanced_search(username, query, filters)
    
    return {
        "query": query,
        "total_results": len(results),
        "results": results,
        "filters": filters
    }

@app.get("/api/search/suggestions")
async def api_search_suggestions(request: Request):
    """API-Endpoint für Suchvorschläge"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    partial = request.query_params.get('q', '').strip()
    
    if len(partial) < 2:
        return {"suggestions": []}
    
    suggestions = get_search_suggestions(username, partial, limit=10)
    
    return {
        "query": partial,
        "suggestions": suggestions
    }

@app.get("/api/search/stats")
async def api_search_statistics(request: Request):
    """API-Endpoint für Such-Statistiken"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    stats = get_search_statistics_cached(username)
    
    return stats

@app.get("/search/export")
async def export_search_results_route(request: Request):
    """Exportiert Suchergebnisse als CSV"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    # Parameter (gleiche wie bei normaler Suche)
    query = request.query_params.get('q', '').strip()
    thread_id = request.query_params.get('thread', None)
    days_back = request.query_params.get('days', None)
    
    if not query:
        return RedirectResponse("/search?error=no_query", status_code=302)
    
    try:
        days_back = int(days_back) if days_back else None
    except ValueError:
        days_back = None
    
    # Suche ausführen (mehr Ergebnisse für Export)
    results = search_chat_messages(
        username=username,
        query=query,
        thread_id=thread_id,
        days_back=days_back,
        limit=500  # Mehr für Export
    )
    
    if not results:
        return RedirectResponse("/search?error=no_results", status_code=302)
    
    # CSV generieren
    csv_data = export_search_results(username, results)
    
    # Als Download zurückgeben
    filename = f"chat_search_{query[:20]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    
    return StreamingResponse(
        io.BytesIO(csv_data.encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

# ──────────────────────────────
# Quick Search Integration in Chat
# ──────────────────────────────

@app.get("/api/chat/quick-search")
async def api_quick_search_in_thread(request: Request):
    """Schnellsuche innerhalb eines Threads"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    query = request.query_params.get('q', '').strip()
    thread_id = request.query_params.get('thread', get_current_thread_id(request))
    
    if not query or len(query) < 2:
        return {"results": []}
    
    # Suche nur in aktuellem Thread
    results = search_chat_messages(
        username=username,
        query=query,
        thread_id=thread_id,
        limit=20
    )
    
    # Vereinfachte Ergebnisse für Quick-Search
    quick_results = []
    for result in results:
        quick_results.append({
            'content': result['highlighted_content'],
            'role': result['role'],
            'timestamp': result['timestamp']
        })
    
    return {
        "query": query,
        "thread_id": thread_id,
        "results": quick_results
    }

# ──────────────────────────────
# Admin Search Functions
# ──────────────────────────────

@app.get("/admin/search/{username}")
async def admin_search_user_messages(request: Request, username: str):
    """Admin kann User-Nachrichten durchsuchen"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    # User-Threads für Admin-Suche
    threads = get_user_threads(username)
    stats = get_search_statistics(username)
    
    return templates.TemplateResponse("admin_search.html", {
        "request": request,
        "target_username": username,
        "threads": threads,
        "stats": stats
    })

@app.get("/api/admin/search/{username}")
async def api_admin_search_messages(request: Request, username: str):
    """API für Admin-Suche in User-Nachrichten"""
    guard = admin_redirect_guard(request)
    if guard:
        return {"error": "Unauthorized"}
    
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    # Gleiche Parameter wie normale Suche
    query = request.query_params.get('q', '').strip()
    thread_id = request.query_params.get('thread', None)
    days_back = request.query_params.get('days', None)
    limit = min(int(request.query_params.get('limit', 50)), 200)  # Admins dürfen mehr
    
    if not query:
        return {"error": "Suchbegriff fehlt"}
    
    try:
        days_back = int(days_back) if days_back else None
    except ValueError:
        days_back = None
    
    # Suche für Admin ausführen
    results = search_chat_messages(
        username=username,  # Target-User
        query=query,
        thread_id=thread_id,
        days_back=days_back,
        limit=limit
    )
    
    return {
        "target_user": username,
        "query": query,
        "total_results": len(results),
        "results": results
    }

# ──────────────────────────────
# Search Integration with Background Tasks
# ──────────────────────────────

async def search_index_maintenance():
    """Background-Task für Such-Index-Wartung"""
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # SQLite FTS (Full-Text Search) Setup falls nicht vorhanden
        cursor.execute("""
            CREATE VIRTUAL TABLE IF NOT EXISTS search_index 
            USING fts5(username, thread_id, role, content, timestamp)
        """)
        
        # Index mit aktuellen Daten füllen falls leer
        cursor.execute("SELECT COUNT(*) FROM search_index")
        index_count = cursor.fetchone()[0]
        
        if index_count == 0:
            logger.info("[SEARCH] Building search index...")
            cursor.execute("""
                INSERT INTO search_index (username, thread_id, role, content, timestamp)
                SELECT username, thread_id, role, content, timestamp 
                FROM chat_history
            """)
            logger.info("[SEARCH] Search index built")
        
        conn.commit()
        conn.close()
        
        return {
            "index_maintained": True,
            "index_entries": index_count
        }
        
    except Exception as e:
        logger.error(f"Search index maintenance error: {e}")
        raise

# Task registrieren
def setup_search_background_tasks():
    """Registriert Search-bezogene Background-Tasks"""
    task_manager.register_task_function(
        "search_index_maintenance",
        search_index_maintenance,
        "Wartung des Such-Index für bessere Performance"
    )
    
    # Einmalig beim Start ausführen
    task_manager.schedule_recurring_task(
        "search_index_maintenance_scheduled",
        search_index_maintenance,
        interval_seconds=86400,  # Täglich
        run_immediately=True
    )

# ──────────────────────────────
# Enhanced Search with FTS
# ──────────────────────────────

def search_with_fts(username: str, query: str, limit: int = 50) -> List[Dict]:
    """
    Erweiterte Suche mit SQLite Full-Text Search (FTS)
    Deutlich schneller für große Datenmengen
    
    Args:
        username: Der User
        query: Suchbegriff
        limit: Max. Anzahl Ergebnisse
    
    Returns:
        Liste von Suchergebnissen
    """
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    
    try:
        # Prüfen ob FTS-Tabelle existiert
        cursor.execute("""
            SELECT name FROM sqlite_master 
            WHERE type='table' AND name='search_index'
        """)
        
        if not cursor.fetchone():
            # Fallback zu normaler Suche
            return search_chat_messages(username, query, limit=limit)
        
        # FTS-Suche
        cursor.execute("""
            SELECT h.id, h.role, h.content, h.timestamp, h.thread_id,
                   t.title as thread_title,
                   snippet(search_index, 2, '<mark>', '</mark>', '...', 30) as snippet
            FROM search_index s
            JOIN chat_history h ON h.username = s.username AND h.content = s.content
            LEFT JOIN chat_threads t ON h.thread_id = t.id AND h.username = t.username
            WHERE s.username = ? AND search_index MATCH ?
            ORDER BY rank
            LIMIT ?
        """, (username, query, limit))
        
        rows = cursor.fetchall()
        results = []
        
        for row in rows:
            context = get_message_context(username, row[0], row[4])
            
            results.append({
                'id': row[0],
                'role': row[1],
                'content': row[2],
                'timestamp': row[3],
                'thread_id': row[4],
                'thread_title': row[5] or 'Unbekannter Thread',
                'highlighted_content': row[6] or highlight_search_term(row[2], query),
                'context': context,
                'search_method': 'fts'
            })
        
        return results
        
    except Exception as e:
        logger.error(f"[SEARCH] FTS error for {username}: {e}")
        # Fallback zu normaler Suche
        return search_chat_messages(username, query, limit=limit)
    finally:
        conn.close()

# ──────────────────────────────
# Search Cache Invalidation
# ──────────────────────────────

def invalidate_search_on_new_message(username: str, thread_id: str):
    """
    Invalidiert Such-Cache wenn neue Nachricht hinzugefügt wird
    Sollte nach save_message_to_thread() aufgerufen werden
    """
    invalidate_search_cache(username)
    
    # FTS-Index aktualisieren falls vorhanden
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # Prüfen ob FTS-Tabelle existiert
        cursor.execute("""
            SELECT name FROM sqlite_master 
            WHERE type='table' AND name='search_index'
        """)
        
        if cursor.fetchone():
            # Neueste Nachricht zum Index hinzufügen
            cursor.execute("""
                INSERT INTO search_index (username, thread_id, role, content, timestamp)
                SELECT username, thread_id, role, content, timestamp 
                FROM chat_history 
                WHERE username = ? AND thread_id = ?
                ORDER BY id DESC 
                LIMIT 1
            """, (username, thread_id))
            
            conn.commit()
        
        conn.close()
        
    except Exception as e:
        logger.error(f"[SEARCH] Error updating FTS index: {e}")
# ──────────────────────────────
# Thread Management Routes
# ──────────────────────────────

@app.post("/threads/new")
async def create_thread(request: Request, title: str = Form(None)):
    """Erstellt einen neuen Thread"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    try:
        thread_id = create_new_thread(username, title)
        set_current_thread_id(request, thread_id)
        invalidate_thread_cache(username)
        
        logger.info(f"[THREADING] User {username} created new thread: {thread_id}")
        return RedirectResponse(f"/chat?thread={thread_id}", status_code=302)
        
    except Exception as e:
        logger.error(f"[THREADING] Error creating thread for {username}: {e}")
        return RedirectResponse("/chat?error=thread_creation_failed", status_code=302)

@app.post("/threads/{thread_id}/rename")
async def rename_thread(request: Request, thread_id: str, new_title: str = Form(...)):
    """Benennt einen Thread um"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    if update_thread_title(username, thread_id, new_title):
        invalidate_thread_cache(username)
        return {"success": True, "message": "Thread umbenannt"}
    else:
        return {"success": False, "message": "Thread konnte nicht umbenannt werden"}

@app.post("/threads/{thread_id}/archive")
async def archive_thread_route(request: Request, thread_id: str):
    """Archiviert einen Thread"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    if archive_thread(username, thread_id):
        invalidate_thread_cache(username)
        
        # Wenn aktueller Thread archiviert wird, zu default wechseln
        if get_current_thread_id(request) == thread_id:
            set_current_thread_id(request, 'default')
        
        return {"success": True, "message": "Thread archiviert"}
    else:
        return {"success": False, "message": "Thread konnte nicht archiviert werden"}

@app.post("/threads/{thread_id}/delete")
async def delete_thread_route(request: Request, thread_id: str):
    """Löscht einen Thread komplett"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    if delete_thread_completely(username, thread_id):
        invalidate_thread_cache(username)
        
        # Wenn aktueller Thread gelöscht wird, zu default wechseln
        if get_current_thread_id(request) == thread_id:
            set_current_thread_id(request, 'default')
        
        return {"success": True, "message": "Thread gelöscht"}
    else:
        return {"success": False, "message": "Thread konnte nicht gelöscht werden"}

@app.get("/api/threads")
async def api_get_threads(request: Request):
    """API-Endpoint für Thread-Liste (AJAX)"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    threads = get_user_threads(username)
    
    return {
        "threads": [
            {
                "id": t['id'],
                "title": t['title'],
                "message_count": t['message_count'],
                "is_archived": t['is_archived'],
                "last_message_preview": t.get('last_message', '')[:50] + "..." if len(t.get('last_message', '')) > 50 else t.get('last_message', ''),
                "updated_at": t['updated_at'].isoformat() if t['updated_at'] else None
            }
            for t in threads
        ],
        "current_thread": get_current_thread_id(request)
    }

@app.get("/api/threads/{thread_id}/history")
async def api_get_thread_history(request: Request, thread_id: str):
    """API-Endpoint für Thread-Historie (AJAX)"""
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    username = request.session.get("username")
    
    # Berechtigung prüfen
    thread_info = get_thread_info(username, thread_id)
    if not thread_info and thread_id != 'default':
        return {"error": "Thread not found"}
    
    history = get_thread_history(username, thread_id)
    
    return {
        "thread_id": thread_id,
        "thread_info": thread_info,
        "history": [
            {
                "role": msg['role'],
                "content": render_markdown_simple(msg['content']),
                "raw_content": msg['content'],
                "timestamp": msg['timestamp']
            }
            for msg in history
        ]
    }

# ──────────────────────────────
# Legacy Route Compatibility
# ──────────────────────────────

@app.post("/chat/clear-history")
async def clear_thread_history(request: Request, thread_id: str = Form(None)):
    """Löscht Thread-Verlauf (kompatibel mit alter Route)"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    current_thread = thread_id or get_current_thread_id(request)
    
    # Thread-Nachrichten löschen
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM chat_history WHERE username = ? AND thread_id = ?", 
                  (username, current_thread))
    
    # Thread-Nachrichten-Counter zurücksetzen
    cursor.execute("UPDATE chat_threads SET message_count = 0, updated_at = ? WHERE id = ? AND username = ?",
                  (datetime.now(), current_thread, username))
    
    conn.commit()
    conn.close()
    
    invalidate_thread_cache(username)
    logger.info(f"[THREADING] Thread history cleared: {current_thread} for {username}")
    
    return RedirectResponse(f"/chat?thread={current_thread}", status_code=302)

# ──────────────────────────────
# Admin Thread Management
# ──────────────────────────────

@app.get("/admin/threads/{username}")
async def admin_view_user_threads(request: Request, username: str):
    """Admin kann User-Threads einsehen"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    threads = get_user_threads(username)
    
    return templates.TemplateResponse("admin_threads.html", {
        "request": request,
        "target_username": username,
        "threads": threads
    })

@app.post("/admin/threads/{username}/{thread_id}/delete")
async def admin_delete_thread(request: Request, username: str, thread_id: str):
    """Admin kann User-Threads löschen"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    if delete_thread_completely(username, thread_id):
        logger.info(f"[ADMIN] Thread deleted by admin: {thread_id} from {username}")
        return {"success": True, "message": "Thread gelöscht"}
    else:
        return {"success": False, "message": "Thread konnte nicht gelöscht werden"}

# ──────────────────────────────
# Routes - Persona
# ──────────────────────────────
@app.get("/persona", response_class=HTMLResponse)
async def persona_settings(request: Request):
    """Persona-Einstellungen Seite mit Subscription-Support"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    current_persona = get_user_persona_cached(username)
    user_tier = get_user_subscription_tier_cached(username)
    
    # Verfügbare Personas für User-Tier
    available_personas = get_available_personas_for_user_cached(username)
    
    # Tier-Informationen
    tier_info = SUBSCRIPTION_TIERS.get(user_tier, SUBSCRIPTION_TIERS["free"])
    
    return templates.TemplateResponse("persona.html", {
        "request": request,
        "username": username,
        "personas": PERSONAS,
        "available_personas": available_personas,
        "current_persona": current_persona,
        "subscription_tier": user_tier,
        "tier_info": tier_info,
        "subscription_tiers": SUBSCRIPTION_TIERS
    })

@app.post("/persona")
async def set_persona(request: Request, persona: str = Form(...)):
    """Persona auswählen mit Subscription-Prüfung"""
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    available_personas = get_available_personas_for_user_cached(username)
    
    if persona in available_personas:
        save_user_persona_cached(username, persona)
        logger.info(f"[PERSONA] {username} wählte Persona: {persona}")
        return RedirectResponse("/chat", status_code=302)
    else:
        # Persona nicht verfügbar für User-Tier
        return templates.TemplateResponse("persona.html", {
            "request": request,
            "username": username,
            "personas": PERSONAS,
            "available_personas": available_personas,
            "current_persona": get_user_persona(username),
            "subscription_tier": get_user_subscription_tier(username),
            "tier_info": SUBSCRIPTION_TIERS.get(get_user_subscription_tier(username), SUBSCRIPTION_TIERS["free"]),
            "subscription_tiers": SUBSCRIPTION_TIERS,
            "error": f"Persona '{PERSONAS.get(persona, {}).get('name', persona)}' ist nicht in deinem aktuellen Plan verfügbar."
        })

# ──────────────────────────────
# Routes - Admin
# ──────────────────────────────
@app.get("/admin", response_class=HTMLResponse)
async def admin_page(request: Request):
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    # Session-Check auch für Admin
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    users = get_all_users_cached()
    return templates.TemplateResponse("admin_users.html", {
        "request": request, 
        "users": users,
        "subscription_tiers": SUBSCRIPTION_TIERS
    })

@app.post("/admin/toggle-block")
async def toggle_block_user(request: Request, username: str = Form(...)):
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    if username != "admin":  # Admin kann sich nicht selbst blockieren
        toggle_user_block_cached(username)
        logger.info(f"[ADMIN] User blockiert/freigeschaltet: {username}")
    
    return RedirectResponse("/admin", status_code=302)

@app.get("/admin/cache-stats")
async def cache_stats(request: Request):
    """Cache-Statistiken für Admin"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    return {
        "cache_stats": app_cache.stats(),
        "cleanup_available": True
    }

@app.post("/admin/clear-cache")
async def clear_cache(request: Request):
    """Cache leeren (Admin)"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    app_cache.clear()
    logger.info("[CACHE] Admin hat Cache geleert")
    return {"message": "Cache geleert", "success": True}

@app.post("/admin/change-tier")
async def change_user_tier(request: Request, 
                          username: str = Form(...),
                          tier: str = Form(...)):
    """Admin kann User-Subscription-Tier ändern"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    if tier in SUBSCRIPTION_TIERS:
        set_user_subscription_tier_cached(username, tier)
        logger.info(f"[ADMIN] Subscription-Tier für {username} geändert zu: {tier}")
    
    return RedirectResponse("/admin", status_code=302)

@app.get("/admin/history/{username}", response_class=HTMLResponse)
async def view_user_history(request: Request, username: str):
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    history = get_user_history(username)
    
    formatted = "<br><br>".join(
        f"<b>{html.escape(str(msg.get('role', 'unknown')))}:</b><br>{html.escape(str(msg.get('content', '')))}"
        for msg in history
    )
    
    body = f"""
    <html><body style="font-family: Arial; padding: 20px;">
    <h1>Chat-Verlauf von {html.escape(username)}</h1>
    <div style="background: #f5f5f5; padding: 15px; border-radius: 5px; margin: 20px 0;">
        {formatted or '— Kein Verlauf vorhanden —'}
    </div>
    <a href='/admin' style="color: blue;">🔙 Zurück zum Admin-Panel</a>
    </body></html>
    """
    return HTMLResponse(body)

@app.post("/admin/delete-history")
async def delete_history(request: Request, username: str = Form(...)):
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    delete_user_history(username)
    logger.info(f"[ADMIN] Chat-Verlauf gelöscht: {username}")
    
    return RedirectResponse("/admin", status_code=302)

@app.post("/admin/delete-user")
async def delete_user(request: Request, username: str = Form(...)):
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    if username != "admin":  # Admin kann sich nicht selbst löschen
        delete_user_completely(username)
        logger.info(f"[ADMIN] User gelöscht: {username}")
    
    return RedirectResponse("/admin", status_code=302)

@app.post("/admin/change-user-password")
async def change_user_password(request: Request, 
                              username: str = Form(...), 
                              new_password: str = Form(...)):
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    if username != "admin":  # Admin-Passwort separat ändern
        reset_password(username, new_password)
        logger.info(f"[ADMIN] Passwort geändert für: {username}")
    
    return RedirectResponse("/admin", status_code=302)

@app.post("/admin/change-password")
async def change_admin_password(request: Request, 
                               old_password: str = Form(...), 
                               new_password: str = Form(...)):
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    admin_user = get_user("admin")
    if admin_user and hmac.compare_digest(admin_user["password"], hash_password(old_password)):
        reset_password("admin", new_password)
        logger.info("[ADMIN] Admin-Passwort geändert")
        return RedirectResponse("/admin", status_code=302)
    else:
        users = get_all_users()
        return templates.TemplateResponse("admin_users.html", {
            "request": request,
            "users": users,
            "subscription_tiers": SUBSCRIPTION_TIERS,
            "error": "Altes Passwort ist falsch"
        })

@app.get("/admin/export-csv")
async def export_csv():
    users = get_all_users()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Benutzername", "Passwort-Hash", "Frage", "Antwort", "Admin", "Blockiert", "Subscription-Tier"])
    
    for name, data in users.items():
        writer.writerow([
            name,
            data.get("password", ""),
            data.get("question", ""),
            data.get("answer", ""),
            "Ja" if data.get("is_admin") else "Nein",
            "Ja" if data.get("blocked") else "Nein",
            data.get("subscription_tier", "free")
        ])
    
    output.seek(0)
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode("utf-8")),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=users.csv"}
    )

@app.get("/admin/performance", response_class=HTMLResponse)
async def admin_performance(request: Request):
    """Admin-Seite für Performance-Monitoring"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    # Session-Check auch für Admin
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    # Performance-Daten vom globalen Monitor holen
    stats = PerformanceMonitoringMiddleware.instance.get_stats()
    cache_stats = app_cache.stats()
    
    return templates.TemplateResponse("admin_performance.html", {
        "request": request,
        "stats": stats,
        "cache_stats": cache_stats
    })

@app.get("/api/performance-stats")
async def api_performance_stats(request: Request):
    """API für Performance-Statistiken (für AJAX-Updates)"""
    guard = admin_redirect_guard(request)
    if guard:
        return {"error": "Unauthorized"}
    
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    return PerformanceMonitoringMiddleware.instance.get_stats()

@app.get("/admin/database", response_class=HTMLResponse)
async def admin_database(request: Request):
    guard = admin_redirect_guard(request)
    if guard:
        return guard

    redirect = require_active_session(request)
    if redirect:
        return redirect

    # --- SQLite-Stats einsammeln ---
    stats = {}
    try:
        size_bytes = os.path.getsize(DB_PATH) if os.path.exists(DB_PATH) else 0
        conn = sqlite3.connect(DB_PATH)
        cur = conn.cursor()

        def fetch_one(sql):
            try:
                cur.execute(sql)
                row = cur.fetchone()
                return row[0] if row else None
            except Exception:
                return None

        # PRAGMAs / Infos
        stats["sqlite_version"]   = sqlite3.sqlite_version
        stats["page_count"]       = fetch_one("PRAGMA page_count;")
        stats["page_size"]        = fetch_one("PRAGMA page_size;")
        stats["freelist_count"]   = fetch_one("PRAGMA freelist_count;")
        stats["journal_mode"]     = fetch_one("PRAGMA journal_mode;")
        stats["wal_autocheckpoint"]= fetch_one("PRAGMA wal_autocheckpoint;")
        stats["cache_size"]       = fetch_one("PRAGMA cache_size;")
        stats["synchronous"]      = fetch_one("PRAGMA synchronous;")
        stats["foreign_keys"]     = fetch_one("PRAGMA foreign_keys;")
        stats["locking_mode"]     = fetch_one("PRAGMA locking_mode;")
        stats["db_size_bytes"]    = size_bytes
        stats["db_size_mb"]       = round((size_bytes or 0) / (1024*1024), 2)

        # Tabellen + Zeilenzahlen (optional)
        try:
            cur.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
            tables = [r[0] for r in cur.fetchall()]
            table_counts = {}
            for t in tables:
                try:
                    cur.execute(f"SELECT COUNT(*) FROM {t}")
                    table_counts[t] = cur.fetchone()[0]
                except Exception:
                    table_counts[t] = None
            stats["tables"] = table_counts
        except Exception:
            stats["tables"] = {}

        conn.close()
    except Exception as e:
        logger.error(f"/admin/database stats error: {e}")

    ctx = {
        "request": request,
        "db_stats": stats,              # <-- WICHTIG
        "db_path": DB_PATH,
    }
    return templates.TemplateResponse("admin_database.html", ctx)

# ──────────────────────────────
# Admin Routes für Background Task Management  
# ──────────────────────────────

@app.get("/admin/tasks", response_class=HTMLResponse)
async def admin_tasks(request: Request):
    """Admin-Seite für Task-Management"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    # Task-Statistiken sammeln
    task_stats = task_manager.get_stats()
    all_tasks = task_manager.get_all_tasks()
    running_tasks = task_manager.get_running_tasks()
    
    # Letzte Tasks (neueste zuerst)
    recent_tasks = sorted(all_tasks, key=lambda x: x.created_at, reverse=True)[:20]
    
    # Registrierte Task-Funktionen
    registered_functions = task_manager.registered_task_functions
    
    # Geplante Tasks
    scheduled_info = []
    for name, schedule in task_manager.scheduled_tasks.items():
        scheduled_info.append({
            'name': name,
            'interval_minutes': round(schedule['interval'] / 60, 1),
            'next_run': schedule['next_run'],
            'last_run': schedule['last_run'],
            'run_count': schedule['run_count']
        })
    
    return templates.TemplateResponse("admin_tasks.html", {
        "request": request,
        "task_stats": task_stats,
        "recent_tasks": recent_tasks,
        "running_tasks": running_tasks,
        "registered_functions": registered_functions,
        "scheduled_tasks": scheduled_info
    })

@app.post("/admin/tasks/run")
async def run_manual_task(request: Request, task_name: str = Form(...)):
    """Startet einen Task manuell"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    if task_name in task_manager.registered_task_functions:
        func_info = task_manager.registered_task_functions[task_name]
        task_id = await task_manager.run_task(
            f"{task_name} (manual)", 
            func_info['function']
        )
        
        logger.info(f"[ADMIN] Manual task started: {task_name} (ID: {task_id})")
        return RedirectResponse(f"/admin/tasks?started={task_id}", status_code=302)
    
    return RedirectResponse("/admin/tasks?error=unknown_task", status_code=302)

@app.post("/admin/tasks/cancel/{task_id}")
async def cancel_task(request: Request, task_id: str):
    """Bricht einen laufenden Task ab"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    success = task_manager.cancel_task(task_id)
    if success:
        logger.info(f"[ADMIN] Task cancelled: {task_id}")
        return {"success": True, "message": "Task abgebrochen"}
    
    return {"success": False, "message": "Task konnte nicht abgebrochen werden"}

@app.get("/api/admin/tasks/status")
async def api_task_status(request: Request):
    """API für Live-Task-Status (AJAX)"""
    guard = admin_redirect_guard(request)
    if guard:
        return {"error": "Unauthorized"}
    
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    stats = task_manager.get_stats()
    running_tasks = [
        {
            "id": task.id,
            "name": task.name,
            "status": task.status.value,
            "progress": task.progress,
            "started_at": task.started_at.isoformat() if task.started_at else None,
            "duration_seconds": (
                (datetime.now() - task.started_at).total_seconds() 
                if task.started_at else 0
            )
        }
        for task in task_manager.get_running_tasks()
    ]
    
    return {
        "stats": stats,
        "running_tasks": running_tasks,
        "timestamp": datetime.now().isoformat()
    }

@app.get("/api/admin/tasks/{task_id}/details")
async def api_task_details(request: Request, task_id: str):
    """Detailinformationen zu einem Task"""
    guard = admin_redirect_guard(request)
    if guard:
        return {"error": "Unauthorized"}
    
    redirect = require_active_session(request)
    if redirect:
        return {"error": "Session expired"}
    
    task_info = task_manager.get_task_info(task_id)
    if not task_info:
        return {"error": "Task not found"}
    
    return {
        "id": task_info.id,
        "name": task_info.name,
        "status": task_info.status.value,
        "created_at": task_info.created_at.isoformat(),
        "started_at": task_info.started_at.isoformat() if task_info.started_at else None,
        "completed_at": task_info.completed_at.isoformat() if task_info.completed_at else None,
        "progress": task_info.progress,
        "error_message": task_info.error_message,
        "result": task_info.result,
        "metadata": task_info.metadata,
        "duration_seconds": (
            (task_info.completed_at - task_info.started_at).total_seconds()
            if task_info.started_at and task_info.completed_at else None
        )
    }

@app.post("/admin/tasks/cleanup")
async def cleanup_old_tasks(request: Request, older_than_hours: int = Form(24)):
    """Bereinigt alte Task-Einträge"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    cleaned_count = task_manager.cleanup_old_tasks(older_than_hours)
    logger.info(f"[ADMIN] Cleaned {cleaned_count} old task entries")
    
    return RedirectResponse(f"/admin/tasks?cleaned={cleaned_count}", status_code=302)

# ──────────────────────────────
# Task-Manager Integration in Startup
# ──────────────────────────────

# Diese Funktion ersetzt deine bisherige startup-Funktion
async def enhanced_startup():
    """Erweiterte Startup-Funktion mit Background-Task-Manager"""
    init_db()
    
    # Background-Task-System initialisieren
    setup_background_tasks()
    
    # Task-Manager-Cleanup-Task starten (ersetzt cache_cleanup_background)
    task_manager.schedule_recurring_task(
        "task_manager_cleanup",
        lambda: task_manager.cleanup_old_tasks(24),
        interval_seconds=21600,  # Alle 6 Stunden
        run_immediately=False
    )
    
    logger.info("[STARTUP] KI-Chat mit erweiterten Background-Tasks gestartet")

# ──────────────────────────────
# Utility-Funktionen für Tasks
# ──────────────────────────────

def get_task_summary() -> dict:
    """Kurze Task-Zusammenfassung für Dashboard"""
    stats = task_manager.get_stats()
    running = task_manager.get_running_tasks()
    
    # Letzter Fehler
    recent_failed = [
        task for task in task_manager.get_all_tasks() 
        if task.status == TaskStatus.FAILED
    ]
    recent_failed.sort(key=lambda x: x.completed_at or x.created_at, reverse=True)
    last_error = recent_failed[0] if recent_failed else None
    
    return {
        'total_tasks': stats['total_tasks'],
        'running_count': stats['running_tasks'],
        'failed_count': stats['failed_tasks'],
        'running_task_names': [task.name for task in running],
        'last_error': {
            'task_name': last_error.name,
            'error_message': last_error.error_message,
            'when': last_error.completed_at.strftime('%H:%M:%S') if last_error.completed_at else 'Unknown'
        } if last_error else None,
        'scheduler_active': stats['scheduler_active']
    }

# API-Endpoint für Dashboard-Widget
@app.get("/api/admin/tasks/summary")
async def api_task_summary(request: Request):
    """Task-Zusammenfassung für Dashboard-Widget"""
    guard = admin_redirect_guard(request)
    if guard:
        return {"error": "Unauthorized"}
    
    return get_task_summary()

@app.get("/admin/analytics", response_class=HTMLResponse)
async def admin_analytics(request: Request):
    """Admin-Analytics Seite"""
    guard = admin_redirect_guard(request)
    if guard:
        return guard
    
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    # Benutzer-Statistiken sammeln
    users = get_all_users()
    total_users = len(users)
    blocked_users = sum(1 for user in users.values() if user.get("blocked", False))
    admin_users = sum(1 for user in users.values() if user.get("is_admin", False))
    
    # Subscription-Tier-Verteilung
    tier_stats = {"free": 0, "pro": 0, "premium": 0}
    for user in users.values():
        tier = user.get("subscription_tier", "free")
        if tier in tier_stats:
            tier_stats[tier] += 1
    
    # Chat-Statistiken (falls verfügbar)
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # Gesamte Chat-Nachrichten
        cursor.execute("SELECT COUNT(*) FROM chat_history")
        total_messages = cursor.fetchone()[0]
        
        # Nachrichten pro Benutzer
        cursor.execute("""
            SELECT username, COUNT(*) as msg_count 
            FROM chat_history 
            GROUP BY username 
            ORDER BY msg_count DESC 
            LIMIT 10
        """)
        top_users = cursor.fetchall()
        
        # Nachrichten der letzten 7 Tage
        cursor.execute("""
            SELECT DATE(timestamp) as date, COUNT(*) as count 
            FROM chat_history 
            WHERE timestamp >= datetime('now', '-7 days')
            GROUP BY DATE(timestamp)
            ORDER BY date DESC
        """)
        daily_stats = cursor.fetchall()
        
        conn.close()
        
    except Exception as e:
        total_messages = 0
        top_users = []
        daily_stats = []
        logger.error(f"Analytics DB error: {e}")
    
    # Performance-Stats (falls Middleware aktiv)
    perf_stats = {}
    if hasattr(PerformanceMonitoringMiddleware, 'instance') and PerformanceMonitoringMiddleware.instance:
        perf_stats = PerformanceMonitoringMiddleware.instance.get_stats()
    
    return templates.TemplateResponse("admin_analytics.html", {
        "request": request,
        "total_users": total_users,
        "blocked_users": blocked_users,
        "admin_users": admin_users,
        "tier_stats": tier_stats,
        "total_messages": total_messages,
        "top_users": top_users,
        "daily_stats": daily_stats,
        "perf_stats": perf_stats,
        "subscription_tiers": SUBSCRIPTION_TIERS
    })

# ──────────────────────────────
# API Routes
# ──────────────────────────────
@app.get("/api/session-info")
async def session_info(request: Request):
    """API für Session-Status"""
    if not request.session.get("username"):
        return {"active": False}
    
    last_activity = request.session.get("last_activity", _pytime.time())
    remaining_seconds = max(0, SESSION_TIMEOUT_SECONDS - (_pytime.time() - last_activity))
    
    return {
        "active": True,
        "username": request.session.get("username"),
        "remaining_minutes": int(remaining_seconds / 60),
        "remaining_seconds": int(remaining_seconds)
    }

# ──────────────────────────────
# Background Tasks
# ──────────────────────────────
async def cache_cleanup_background():
    """Background-Task für Cache-Bereinigung"""
    await asyncio.sleep(60)  # Warten bis App vollständig gestartet
    
    while True:
        try:
            cleaned = app_cache.cleanup_expired()
            if cleaned > 0:
                logger.info(f"[CACHE] {cleaned} abgelaufene Einträge bereinigt")
            await asyncio.sleep(300)  # Alle 5 Minuten
        except Exception as e:
            logger.error(f"[CACHE ERROR] {e}")
            await asyncio.sleep(300)

# ──────────────────────────────
# Startup / Shutdown (mit Files-Upgrade & File-Tasks)
# ──────────────────────────────

def _ensure_runtime_files():
    """Erstellt fehlende Dateien/Ordner und prüft Basisvoraussetzungen."""
    try:
        if RATE_LIMIT_FILE and not os.path.exists(RATE_LIMIT_FILE):
            with open(RATE_LIMIT_FILE, "w", encoding="utf-8") as f:
                json.dump({}, f)
            logger.info(f"[STARTUP] Created {RATE_LIMIT_FILE}")
    except Exception as e:
        logger.error(f"[STARTUP] Could not prepare {RATE_LIMIT_FILE}: {e}")

def _maybe_print_route_map():
    """Optional: Routenübersicht ausgeben (aktivieren mit env PRINT_ROUTE_MAP=1)."""
    if os.getenv("PRINT_ROUTE_MAP", "0") != "1":
        return
    try:
        print("\n=== ROUTE MAP (order matters) ===")
        for r in app.router.routes:
            path = getattr(r, "path", None)
            methods = getattr(r, "methods", None)
            ep = getattr(r, "endpoint", None)
            mod = ep.__module__ if ep else "?"
            nm = ep.__name__ if ep else "?"
            if path:
                print(f"{methods} {path} -> {mod}.{nm}")
        print("=== END ROUTE MAP ===\n")
    except Exception as e:
        logger.error(f"[STARTUP] Route-map print failed: {e}")

@app.on_event("startup")
async def startup():
    """App-Startup: DB initialisieren, Upgrades, Caches, Background-Tasks."""
    try:
        logger.info("[STARTUP] Initializing application…")

        # 1) Dateien/Umgebung vorbereiten
        _ensure_runtime_files()

        # 2) DB initialisieren & Upgrades (Reihenfolge!):
        #    - Basistabellen
        #    - Threading (chat_threads / thread_id)
        #    - Multi-Model (preferred_model, model_usage-Spalten)
        #    - Files (DEIN neues Upgrade)
        init_db()
        upgrade_database_for_threading()
        upgrade_database_for_models()

        # Files-Upgrade nur ausführen, wenn es existiert
        if "upgrade_database_for_files" in globals() and callable(globals()["upgrade_database_for_files"]):
            try:
                upgrade_database_for_files()
                logger.info("[FILES] Database upgraded for files ✅")
            except Exception as e:
                logger.exception(f"[FILES] Database upgrade failed: {e}")
        else:
            logger.info("[FILES] upgrade_database_for_files() not defined — skipped")

        # 3) Cache einmal leeren (sauberer Start)
        if hasattr(app_cache, "clear"):
            app_cache.clear()

        # 4) Background-Tasks registrieren:
        setup_background_tasks()  # dein allgemeiner Task-Manager

        # File-Tasks nur starten, wenn Funktion existiert
        if "setup_file_background_tasks" in globals() and callable(globals()["setup_file_background_tasks"]):
            try:
                setup_file_background_tasks()
                logger.info("[FILES] File background tasks set up ✅")
            except Exception as e:
                logger.exception(f"[FILES] setup_file_background_tasks failed: {e}")
        else:
            logger.info("[FILES] setup_file_background_tasks() not defined — skipped")

        # 5) Optional: Routenliste
        _maybe_print_route_map()

        logger.info("[STARTUP] Application started ✅")
    except Exception as e:
        logger.exception(f"[STARTUP] Failed to start application: {e}")
        raise

@app.on_event("shutdown")
async def shutdown():
    """Sauberes Herunterfahren (Scheduler stoppen, Ressourcen freigeben)."""
    try:
        # Allgemeinen Scheduler stoppen
        if hasattr(task_manager, "scheduler_running"):
            task_manager.scheduler_running = False
        if hasattr(task_manager, "running_tasks"):
            for task_id, a_task in list(task_manager.running_tasks.items()):
                try:
                    a_task.cancel()
                except Exception:
                    pass

        # Falls du einen separaten File-Task-Manager verwendest (z. B. file_task_manager)
        if "file_task_manager" in globals():
            ftm = globals().get("file_task_manager")
            try:
                if hasattr(ftm, "scheduler_running"):
                    ftm.scheduler_running = False
                if hasattr(ftm, "running_tasks"):
                    for task_id, a_task in list(ftm.running_tasks.items()):
                        try:
                            a_task.cancel()
                        except Exception:
                            pass
            except Exception as e:
                logger.error(f"[SHUTDOWN] Error stopping file_task_manager: {e}")

        logger.info("[SHUTDOWN] Application shutdown complete ✅")
    except Exception as e:
        logger.error(f"[SHUTDOWN] Error during shutdown: {e}")

from services.shopping import generate_shopping_list
from services.pricing import calculate_pricing
from services.recipes import generate_recipes
from services.exports import generate_csv_shopping_list, generate_csv_pricing

# EINKAUFSLISTE
@app.get("/gastro/shopping-list", response_class=HTMLResponse)
async def shopping_list_page(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    
    return templates.TemplateResponse("gastro/shopping_list.html", {
        "request": request,
        "username": username
    })

@app.post("/gastro/shopping-list")
async def shopping_list_generate(
    request: Request,
    gericht: str = Form(...),
    portionen: int = Form(...)
):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    result = generate_shopping_list(gericht, portionen)
    
    request.session["last_shopping_result"] = result
    
    return templates.TemplateResponse("gastro/shopping_list.html", {
        "request": request,
        "username": username,
        "result": result
    })

@app.get("/gastro/shopping-list/export/csv")
async def export_shopping_csv(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    result = request.session.get("last_shopping_result")
    if not result:
        return RedirectResponse("/gastro/shopping-list")
    
    csv_content = generate_csv_shopping_list(result)
    
    return StreamingResponse(
        io.BytesIO(csv_content.encode('utf-8')),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=einkaufsliste.csv"}
    )

# PREISKALKULATION
@app.get("/gastro/pricing", response_class=HTMLResponse)
async def pricing_page(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    return templates.TemplateResponse("gastro/pricing.html", {
        "request": request,
        "username": username
    })

@app.post("/gastro/pricing")
async def pricing_calculate(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    form_data = await request.form()
    
    gericht = form_data.get("gericht")
    zielmarge = float(form_data.get("zielmarge", 250))
    
    zutaten = []
    i = 0
    while f"zutat_name_{i}" in form_data:
        zutaten.append({
            "name": form_data.get(f"zutat_name_{i}"),
            "menge": float(form_data.get(f"zutat_menge_{i}", 1)),
            "ek_preis": float(form_data.get(f"zutat_preis_{i}", 0))
        })
        i += 1
    
    result = calculate_pricing(zutaten, zielmarge)
    request.session["last_pricing_result"] = {**result, "gericht": gericht}
    
    return templates.TemplateResponse("gastro/pricing.html", {
        "request": request,
        "username": username,
        "result": result,
        "gericht": gericht
    })

@app.get("/gastro/pricing/export/csv")
async def export_pricing_csv(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    result = request.session.get("last_pricing_result")
    if not result:
        return RedirectResponse("/gastro/pricing")
    
    gericht = result.pop("gericht", "Unbenannt")
    csv_content = generate_csv_pricing(result, gericht)
    
    return StreamingResponse(
        io.BytesIO(csv_content.encode('utf-8')),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=kalkulation.csv"}
    )

# REZEPTE
@app.get("/gastro/recipes", response_class=HTMLResponse)
async def recipes_page(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    return templates.TemplateResponse("gastro/recipes.html", {
        "request": request,
        "username": username
    })

@app.post("/gastro/recipes")
async def recipes_generate(
    request: Request,
    zutaten: str = Form(...),
    vegetarisch: bool = Form(False),
    vegan: bool = Form(False)
):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    recipes = generate_recipes(zutaten, vegetarisch, vegan)
    
    return templates.TemplateResponse("gastro/recipes.html", {
        "request": request,
        "username": username,
        "recipes": recipes
    })        
# PRICING
@app.get("/gastro/pricing", response_class=HTMLResponse)
async def pricing_page(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    return templates.TemplateResponse("gastro/pricing.html", {
        "request": request,
        "username": username
    })

@app.post("/gastro/pricing")
async def pricing_calculate(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    form_data = await request.form()
    
    gericht = form_data.get("gericht")
    zielmarge = float(form_data.get("zielmarge", 250))
    
    # Parse Zutaten
    zutaten = []
    i = 0
    while f"zutat_name_{i}" in form_data:
        zutaten.append({
            "name": form_data.get(f"zutat_name_{i}"),
            "menge": float(form_data.get(f"zutat_menge_{i}", 1)),
            "ek_preis": float(form_data.get(f"zutat_preis_{i}", 0))
        })
        i += 1
    
    result = calculate_pricing(zutaten, zielmarge)
    
    # Speichere für CSV
    request.session["last_pricing_result"] = {**result, "gericht": gericht}
    
    return templates.TemplateResponse("gastro/pricing.html", {
        "request": request,
        "username": username,
        "result": result,
        "gericht": gericht
    })

@app.get("/gastro/pricing/export/csv")
async def export_pricing_csv(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    result = request.session.get("last_pricing_result")
    if not result:
        return RedirectResponse("/gastro/pricing")
    
    gericht = result.pop("gericht", "Unbenannt")
    csv_content = generate_csv_pricing(result, gericht)
    
    return StreamingResponse(
        io.BytesIO(csv_content.encode('utf-8')),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=kalkulation.csv"}
    )

# RECIPES
@app.get("/gastro/recipes", response_class=HTMLResponse)
async def recipes_page(request: Request):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    return templates.TemplateResponse("gastro/recipes.html", {
        "request": request,
        "username": username
    })

@app.post("/gastro/recipes")
async def recipes_generate(
    request: Request,
    zutaten: str = Form(...),
    vegetarisch: bool = Form(False),
    vegan: bool = Form(False)
):
    redirect = require_active_session(request)
    if redirect:
        return redirect
    
    username = request.session.get("username")
    recipes = generate_recipes(zutaten, vegetarisch, vegan)
    
    return templates.TemplateResponse("gastro/recipes.html", {
        "request": request,
        "username": username,
        "recipes": recipes
    })
